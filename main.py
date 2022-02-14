from enum import auto
import hashlib
from inspect import CO_NESTED
import urllib
from api_logic import train_face,check_confidence,execute_request, get_name
import logging
from PIL import Image
#from flask_limiter import Limiter
#from flask_limiter.util import get_remote_address
from multiprocessing import connection
import ssl, csv
from tkinter import DOTBOX
from typing import ContextManager
import pyqrcode
from docx import Document
from docxcompose.composer import Composer
from email.mime.image import MIMEImage
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from flask_mobility import Mobility
import os
from datetime import *
from flask import Flask, request, render_template, g, redirect, url_for, flash, session,send_from_directory, send_file
import pyotp
import os
import smtplib
import hashlib
import re
import random
import bcrypt
from forms import FileSubmit, Patient_Login_form, Admin_Login_form,Otp, Register, RequestPatientInfo_Form, Appointment, \
    RegisterDoctor, RegisterResearcher, RegisterHr,General_UpdateForm, Assign_PhysiciantForm, Baseinfo
from functools import wraps
import pyodbc
import textwrap
from customlogging import converTxtToCSV, patientFileModificationFilter,db_log,lr_log,ac_log,virus_log
from mssql_auth import database, server,backup_server

from flask_qrcode import QRcode
from datetime import datetime
from Google import MyDrive
from flask_apscheduler import APScheduler
import atexit
import time
from virustotal import virusTotal
import os.path,base64
from virustotal_python import Virustotal
from pprint import pprint # pprint is used to pretty print in good json format instead of in a line
from ps import psrun
context = ssl.create_default_context()

salt = bcrypt.gensalt() 
# db_connection = pyodbc.connect( # 
# 'DRIVER={ODBC Driver 17 for SQL Server}; \
# SERVER=' + server+ '; \
# DATABASE=' + 'database1' + ';\
# Trusted_Connection=yes; Encrypt=yes;TrustServerCertificate=yes',autocommit=True)



patientFilesModificationLogger=logging.getLogger(__name__+"patientFileModification")
patientFilesModificationLogger.setLevel(logging.DEBUG)
formatterserialize=logging.Formatter('%(asctime)s;%(levelname)s;%(message)s;%(ipaddress)s;%(username)s')
file_handlerModification=logging.FileHandler('logs/patientFileChangelog.txt')
file_handlerModification.setFormatter(formatterserialize)
patientFilesModificationLogger.addHandler(file_handlerModification)



db_logLogger=logging.getLogger(__name__+"db_log")
db_logLogger.setLevel(logging.CRITICAL)
db_formatter=logging.Formatter('%(asctime)s;%(levelname)s;%(message)s')
db_logLoggerModification=logging.FileHandler('logs/db_logs.txt')
db_logLoggerModification.setFormatter(db_formatter)
db_logLogger.addHandler(db_logLoggerModification)

lr_logLogger=logging.getLogger(__name__+"lr_log")
lr_logLogger.setLevel(logging.DEBUG)
lr_formatter=logging.Formatter('%(asctime)s;%(levelname)s;%(username)s;%(message)s')
lr_logLoggerModification=logging.FileHandler('logs/lr_logs.txt')
lr_logLoggerModification.setFormatter(lr_formatter)
lr_logLogger.addHandler(lr_logLoggerModification)


ac_logLogger=logging.getLogger(__name__+"ac_log")
ac_logLogger.setLevel(logging.WARNING)
ac_formatter=logging.Formatter('%(asctime)s;%(levelname)s;%(message)s;%(ipaddress)s;%(username)s;%(page_visited)s')
ac_logLoggerModification=logging.FileHandler('logs/ac_logs.txt')
ac_logLoggerModification.setFormatter(ac_formatter)
ac_logLogger.addHandler(ac_logLoggerModification)

virus_logLogger=logging.getLogger(__name__+"virus_log")
virus_logLogger.setLevel(logging.CRITICAL)
virus_formatter=logging.Formatter('%(asctime)s;%(levelname)s;%(message)s;%(ipaddress)s;%(username)s')
virus_logLoggerModification=logging.FileHandler('logs/virus_logs.txt')
virus_logLoggerModification.setFormatter(virus_formatter)
virus_logLogger.addHandler(virus_logLoggerModification)




# def autonomous_backup():
#     auto_use_seconddb()
#     # t = time.localtime()
#     # current_time = str(time.strftime("%d %B %Y_%H;%M;%S",t)) #use semicolon cuz window does not allow colon

#     # backup = f"BACKUP DATABASE [database1] TO DISK = N'C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\{current_time}.bak'"
#     # cur = connection.cursor()
#     # cur.execute(backup)
#     # while (cur.nextset()):n
#     #     pass
#     # print('Local Backup successful')

#     # Update to Google Drive
#     folder_path = "C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\"
#     folders = os.listdir(folder_path) #['folder1', 'folder2'] list folder

#     my_drive.create_file(folders[-1], folder_path) #function in backup.py file
#     print('Drive backup successful')
#     print("autonomous backup")

#     scheduler.remove_all_jobs()
#     # time.sleep(15)
#     scheduler.add_job(id = 'Scheduled Task',func = autonomous_backup, trigger="interval", seconds=10)

#     # flash('Local and cloud backup successful')


def autonomous_conversion():
    converTxtToCSV()
    scheduler.remove_all_jobs()
    scheduler.add_job(id = 'Scheduled Task',func = autonomous_conversion, trigger="interval", minutes=30)




app = Flask(__name__)
QRcode(app)
app.config['SECRET_KEY'] = "secret key"
app.config['UPLOAD_FOLDER'] = os.path.join(os.getcwd(),'saved')
app.config['previousState']="None"

def custom_login_required(f):
    @wraps(f)
    def wrap(*args,**kwargs):
        if session is None:
            flash("Please log in to access this page","error")
            print(request.referrer)
            filter = ac_log(request.referrer,request.remote_addr,"")
            ac_logLogger.addFilter(filter)
            ac_logLogger.warning(msg=f"code 401")
            return redirect(url_for('login'))

        if 'login' not in session or session['login']!=True:
            print(request.referrer)
            flash("Please log in to access this page","error")
            filter = ac_log(request.referrer,request.remote_addr,"")
            ac_logLogger.addFilter(filter)
            ac_logLogger.warning(msg=f"code 401")
            return redirect(url_for('login'))


        print(session['login'],'wrapper session logged in')

        return f(*args,**kwargs)

    return wrap



# def auto_use_seconddb():
#
#         x=psrun("Get-Service 'MSSQLSERVER' | Select-Object -Property Status").stdout.read()
#         print(x)
#         db_dict = {}
#         #Primary Server not runningre
#         if b'Stopped' in x:
#             try:
#                 session['db'] = 'Backup Server'
#             except:
#                 pass
#             t = time.localtime()
#
#             current_time = str(time.strftime("%d %B %Y_%H;%M;%S",t)) #use semicolon cuz window does not allow colon
#
#             try:# If unable to connect to second server, restore the second server DB using the bak file from first server db and use the second server connection immediately, as we may not know if first server db is compromised
#                 db_connection = pyodbc.connect( #
#                 'DRIVER={ODBC Driver 17 for SQL Server}; \
#                 SERVER=' + backup_server+ '; \
#                 DATABASE=' + 'database1' + ';\
#                 Trusted_Connection=yes; Encrypt=yes;TrustServerCertificate=yes',autocommit=True)
#                 cur = db_connection.cursor()
#                 if app.config['previousState']=="secondary":
#                     backup = f"BACKUP DATABASE [database1] TO DISK = N'C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\{current_time}.bak'"
#                     cur = db_connection.cursor()
#                     cur.execute(backup)
#                     while (cur.nextset()):
#                         pass
#                     print('Local Backup successful')
#                     folder_path = "C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\"
#                     folders = os.listdir(folder_path) #['folder1', 'folder2'] list folder
#                     my_drive.create_file(folders[-1], folder_path) #function in backup.py file
#                     filter = db_log('primary')
#                     db_logLogger.addFilter(filter)
#                     db_logLogger.critical(
#                         msg=f"secondary server was successfully backed up"
#                     )
#
#                 db_connection = pyodbc.connect( #
#                 'DRIVER={ODBC Driver 17 for SQL Server}; \
#                 SERVER=' + backup_server+ '; \
#                 DATABASE=' + 'database1' + ';\
#                 Trusted_Connection=yes; Encrypt=yes;TrustServerCertificate=yes',autocommit=True)
#                 print("Update and restore secondary db ")
#                 print("Using second db!")
#                 executelock = ("alter database database1 set offline with rollback immediate ")
#                 releaselock = ("alter database database1 set online")
#                 folder_path = "C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\"
#                 folders = os.listdir(folder_path) #['folder1', 'folder2'] list folder
#                 statement = (
#                     f"RESTORE DATABASE database1 FROM  DISK = 'C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\{folders[-1]}'  WITH RECOVERY, MOVE 'database' TO 'C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\MSSQL15.MSSQLSERVER01\\MSSQL\DATA\\database.mdf',MOVE 'database_log' TO 'C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\MSSQL15.MSSQLSERVER01\\MSSQL\DATA\\database_log.ldf', REPLACE" ) #  file name and db must be the same
#                 create_db = ("IF NOT EXISTS( SELECT * FROM sys.databases WHERE name = 'database1' ) BEGIN CREATE DATABASE database1 ; END;")
#                 print(cur.execute(create_db),'createdb')
#                 cur.execute(create_db)
#                 cur.execute(executelock)
#                 cur.execute(statement)
#                 while cur.nextset():
#                     pass
#                 cur.execute(releaselock)
#                 app.config['previousState']="secondary"
#                 filter = db_log(app.config['previousState'])
#                 db_logLogger.addFilter(filter)
#                 db_logLogger.critical(
#                     msg=f"primary server is down"
#                 )
#                 db_logLogger.critical(
#                     msg=f"connection to {app.config['previousState']} server successful"
#                 )
#
#                 return db_connection
#
#             except: #if secondary db was deleted, it will auto restore
#                 db_connection = pyodbc.connect( #
#                 'DRIVER={ODBC Driver 17 for SQL Server}; \
#                 SERVER=' + backup_server+ '; \
#                 Trusted_Connection=yes; Encrypt=yes;TrustServerCertificate=yes',autocommit=True)
#                 cur = db_connection.cursor()
#
#                 executelock = ("alter database database1 set offline with rollback immediate ")
#                 releaselock = ("alter database database1 set online")
#                 folder_path = "C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\"
#                 folders = os.listdir(folder_path) #['folder1', 'folder2'] list folder
#                 if len(folders)==0:
#                     try:
#                         backup = f"BACKUP DATABASE [database1] TO DISK = N'C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\{current_time}.bak'"
#                         cur = db_connection.cursor()
#                         cur.execute(backup)
#                         while (cur.nextset()):
#                             pass
#                         print('Local Backup successful')
#                         folder_path = "C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\"
#                         folders = os.listdir(folder_path) #['folder1', 'folder2'] list folder
#
#                         my_drive.create_file(folders[-1], folder_path) #function in backup.py file
#                     except:
#                         last_file= my_drive.last_file()
#                         my_drive.download_file(last_file)
#                         folders = os.listdir(folder_path) #['folder1', 'folder2'] list folder
#                         filter = db_log('secondary')
#                         db_logLogger.addFilter(filter)
#                         db_logLogger.critical(
#                                 msg=f"there is no backup files in local storage"
#                             )
#                         db_logLogger.critical(
#                             msg=f"pulling backup files from cloud storage"
#                         )
#
#                 statement = (
#                         f"RESTORE DATABASE database1 FROM  DISK = 'C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\{folders[-1]}'  WITH RECOVERY, MOVE 'database' TO 'C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\MSSQL15.MSSQLSERVER01\\MSSQL\DATA\\database.mdf',MOVE 'database_log' TO 'C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\MSSQL15.MSSQLSERVER01\\MSSQL\DATA\\database_log.ldf', REPLACE" ) #  file name and db must be the same
#                 create_db = ("IF NOT EXISTS( SELECT * FROM sys.databases WHERE name = 'database1' ) BEGIN CREATE DATABASE database1 ; END;")
#                 print(cur.execute(create_db),'createdb')
#                 cur.execute(create_db)
#                 cur.execute(executelock)
#                 cur.execute(statement)
#                 while cur.nextset():
#                     pass
#                 cur.execute(releaselock)
#                 backup = f"BACKUP DATABASE [database1] TO DISK = N'C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\{current_time}.bak'"
#                 cur = db_connection.cursor()
#                 cur.execute(backup)
#                 while (cur.nextset()):
#                     pass
#                 folder_path = "C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\"
#                 folders = os.listdir(folder_path) #['folder1', 'folder2'] list folder
#
#                 my_drive.create_file(folders[-1], folder_path) #function in backup.py file
#                 print('Local Backup and drive successful')
#
#                 db_connection = pyodbc.connect( #
#                             'DRIVER={ODBC Driver 17 for SQL Server}; \
#                             SERVER=' + backup_server+ '; \
#                             DATABASE=' + 'database1' + ';\
#                             Trusted_Connection=yes; Encrypt=yes;TrustServerCertificate=yes',autocommit=True)
#
#                 print("Restored second db as it was never created.")
#                 print("Using second db!")
#                 print(';')
#                 app.config['previousState']="secondary"
#                 filter = db_log(app.config['previousState'])
#                 db_logLogger.addFilter(filter)
#                 db_logLogger.critical(
#                     msg=f"{app.config['previousState']} server was deleted"
#             )
#                 db_logLogger.critical(
#                     msg=f"{app.config['previousState']} server database was restored"
#                 )
#                 db_logLogger.critical(
#                     msg=f"connection to {app.config['previousState']} server successful"
#                 )
#
#
#
#                 return db_connection
#
#
#         else:
#             try:
#                 session['db'] = 'Primary Server'
#             except:
#                 pass
#             t = time.localtime()
#             current_time = str(time.strftime("%d %B %Y_%H;%M;%S",t)) #use semicolon cuz window does not allow colon
#
#             try: # If unable to connect to primary server, restore the primary server DB using the bak file from first server db and use the second server connection immediately, as we may not know if first server db is compromised
#                 db_connection = pyodbc.connect( #
#                 'DRIVER={ODBC Driver 17 for SQL Server}; \
#                 SERVER=' + server+ '; \
#                 DATABASE=' + 'database1' + ';\
#                 Trusted_Connection=yes; Encrypt=yes;TrustServerCertificate=yes',autocommit=True)
#                 cur = db_connection.cursor()
#
#                 if app.config['previousState']=="primary" or app.config['previousState']=="None" :
#                     backup = f"BACKUP DATABASE [database1] TO DISK = N'C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\{current_time}.bak'"
#                     cur = db_connection.cursor()
#                     cur.execute(backup)
#                     while (cur.nextset()):
#                         pass
#                     print('Local Backup successful')
#                     folder_path = "C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\"
#                     folders = os.listdir(folder_path) #['folder1', 'folder2'] list folder
#
#                     my_drive.create_file(folders[-1], folder_path) #function in backup.py file
#                     filter = db_log('primary')
#                     db_logLogger.addFilter(filter)
#                     db_logLogger.critical(
#                         msg=f"primary server was successfully backed up"
#                     )
#
#
#
#                 executelock = ("alter database database1 set offline with rollback immediate ")
#                 releaselock = ("alter database database1 set online")
#                 folder_path = "C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\"
#                 folders = os.listdir(folder_path) #['folder1', 'folder2'] list folder
#                 # if len(folders)==0:
#                 #     backup = f"BACKUP DATABASE [database1] TO DISK = N'C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\{current_time}.bak'"
#                 #     cur = db_connection.cursor()
#                 #     cur.execute(backup)
#                 #     while (cur.nextset()):
#                 #         pass
#                 #     print('Local Backup successful')
#
#                 statement = (f"RESTORE DATABASE database1 FROM  DISK = 'C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\{folders[-1]}'  WITH RECOVERY, MOVE 'database' TO 'C:\\Program Files\\Microsoft SQL Server\\MSSQL15.MSSQLSERVER\\MSSQL\\DATA\\database.mdf',MOVE 'database_log' TO 'C:\\Program Files\\Microsoft SQL Server\\MSSQL15.MSSQLSERVER\\MSSQL\\DATA\\database_log.ldf', REPLACE" ) #  file name and db must be the same
#                 create_db = ("IF NOT EXISTS( SELECT * FROM sys.databases WHERE name = 'database1' ) BEGIN CREATE DATABASE database1 ; END;")
#                 print(cur.execute(create_db),'createdb')
#                 cur.execute(create_db)
#                 cur.execute(executelock)
#                 cur.execute(statement)
#                 while cur.nextset():
#                     pass
#                 cur.execute(releaselock)
#
#
#
#                 db_connection = pyodbc.connect( #
#                 'DRIVER={ODBC Driver 17 for SQL Server}; \
#                 SERVER=' + server+ '; \
#                 DATABASE=' + 'database1' + ';\
#                 Trusted_Connection=yes; Encrypt=yes;TrustServerCertificate=yes',autocommit=True)
#                 print("Update and restore primary db ")
#                 print("Using primary db!")
#                 print(db_connection,'db')
#                 app.config['previousState']="primary"
#                 filter = db_log(app.config['previousState'])
#                 db_logLogger.addFilter(filter)
#                 db_logLogger.critical(
#                     msg=f"connection to {app.config['previousState']} server successful"
#                 )
#                 return db_connection
#
#             except:
#                     db_connection = pyodbc.connect( #
#                     'DRIVER={ODBC Driver 17 for SQL Server}; \
#                     SERVER=' + server+ '; \
#                     DATABASE=' + 'master' + ';\
#                     Trusted_Connection=yes; Encrypt=yes;TrustServerCertificate=yes',autocommit=True)
#                     cur = db_connection.cursor()
#                     # backup = f"BACKUP DATABASE [database1] TO DISK = N'C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\{current_time}.bak'"
#                     # cur.execute(backup)
#                     # while (cur.nextset()):
#                     #     pass
#                     # print('Local Backup successful')
#
#
#                     executelock = ("alter database database1 set offline with rollback immediate ")
#                     releaselock = ("alter database database1 set online")
#                     folder_path = "C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\"
#                     folders = os.listdir(folder_path) #['folder1', 'folder2'] list folder
#
#                     if len(folders)==0:
#                         try:
#                             backup = f"BACKUP DATABASE [database1] TO DISK = N'C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\{current_time}.bak'"
#                             cur = db_connection.cursor()
#                             cur.execute(backup)
#                             while (cur.nextset()):
#                                 pass
#                             print('Local Backup successful')
#                             folder_path = "C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\"
#                             folders = os.listdir(folder_path) #['folder1', 'folder2'] list folder
#
#                             my_drive.create_file(folders[-1], folder_path) #function in backup.py file
#                         except:
#                             last_file= my_drive.last_file()
#                             my_drive.download_file(last_file)
#                             folders = os.listdir(folder_path) #['folder1', 'folder2'] list folder
#                             filter = db_log('primary')
#                             db_logLogger.addFilter(filter)
#                             db_logLogger.critical(
#                                 msg=f"there is no backup files in local storage"
#                             )
#                             db_logLogger.critical(
#                                 msg=f"pulling backup files from cloud storage"
#                             )
#
#                     statement = (
#                             f"RESTORE DATABASE database1 FROM  DISK = 'C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\{folders[-1]}'  WITH RECOVERY, MOVE 'database' TO 'C:\\Program Files\\Microsoft SQL Server\\MSSQL15.MSSQLSERVER\\MSSQL\\DATA\\database.mdf',MOVE 'database_log' TO 'C:\\Program Files\\Microsoft SQL Server\\MSSQL15.MSSQLSERVER\\MSSQL\\DATA\\database_log.ldf', REPLACE" ) #  file name and db must be the same
#                     create_db = ("IF NOT EXISTS( SELECT * FROM sys.databases WHERE name = 'database1' ) BEGIN CREATE DATABASE database1 ; END;")
#                     cur.execute(create_db)
#                     cur.execute(executelock)
#                     cur.execute(statement)
#                     while cur.nextset():
#                         pass
#                     cur.execute(releaselock)
#
#                     db_connection = pyodbc.connect( #
#                                 'DRIVER={ODBC Driver 17 for SQL Server}; \
#                                 SERVER=' + server+ '; \
#                                 DATABASE=' + 'database1' + ';\
#                                 Trusted_Connection=yes; Encrypt=yes; TrustServerCertificate=yes',autocommit=True)
#                     print("Restored primary db as it was never created.")
#
#
#                     print("Using primary db!")
#                     print(';')
#                     app.config['previousState']="primary"
#                     filter = db_log(app.config['previousState'])
#                     db_logLogger.addFilter(filter)
#                     db_logLogger.critical(
#                         msg=f"{app.config['previousState']} server was deleted"
#                     )
#                     db_logLogger.critical(
#                         msg=f"{app.config['previousState']} server database was restored"
#                     )
#                     db_logLogger.critical(
#                         msg=f"connection to {app.config['previousState']} server successful"
#                     )
#
#                     return db_connection

def auto_use_seconddb():
    try:  # Try the first server if connection can be established
        db_connection = pyodbc.connect(
            'DRIVER={ODBC Driver 17 for SQL Server}; \
            SERVER=' + 'DESKTOP-DMO8V7F\MSSQLSERVER01' + '; \
            DATABASE=' + 'database1' + ';\
            Trusted_Connection=yes;'
            , autocommit=True
        )
        print('Using first server db!')
        cur = db_connection.cursor()

    except:  # If unable to connect to second server, restore the second server DB using the bak file from first server db and use the second server connection immediately, as we may not know if first server db is compromised
        try:
            db_connection = pyodbc.connect(  #
                'DRIVER={ODBC Driver 17 for SQL Server}; \
                SERVER=' + 'DESKTOP-DMO8V7F\MSSQLSERVER01' + '; \
                DATABASE=' + 'database1' + ';\
                Trusted_Connection=yes;' \
                , autocommit=True)
            cur = db_connection.cursor()
            should_backup = True
            folder_path = "C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\"
            folders = os.listdir(folder_path)  # ['folder1', 'folder2'] list folder
            drive_files = my_drive.list_files()

            # for files in drive_files:
            get_dbfiles = cur.execute(
                'SELECT folder_id from google ORDER BY backup_date DESC').fetchval()  # get latest file
            print(drive_files, 'drivefiles')
            print(get_dbfiles.strip(), 'latest file')

            if get_dbfiles.strip() in drive_files:
                print(get_dbfiles.strip(), 'not working')
                should_backup = False

            if should_backup == False:  # don't need to backup as it is already backed up.
                db_connection = pyodbc.connect(  # need use back the recovered db
                    'DRIVER={ODBC Driver 17 for SQL Server}; \
                    SERVER=' + 'DESKTOP-75MSPGF' + '; \
                    DATABASE=' + 'database1' + ';\
                    Trusted_Connection=yes;' \
                    , autocommit=True)
                print('DB had been backed up.')
                print("Using second db")
            else:
                executelock = ("alter database database1 set offline with rollback immediate ")
                releaselock = ("alter database database1 set online")
                folder_path = "C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\"
                folders = os.listdir(folder_path)  # ['folder1', 'folder2'] list folder
                statement = (
                    f"RESTORE DATABASE database1 FROM  DISK = 'C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\{folders[-1]}' WITH RECOVERY, MOVE 'database' TO 'C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\MSSQL15.SQLEXPRESS01\\MSSQL\DATA\\database.mdf',MOVE 'database_log' TO 'C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\MSSQL15.SQLEXPRESS01\\MSSQL\DATA\\database_log.ldf', REPLACE")  # file name and db must be the same
                create_db = (
                    "IF NOT EXISTS( SELECT * FROM sys.databases WHERE name = 'database1' ) BEGIN CREATE DATABASE database1 ; END;")
                print(cur.execute(create_db), 'createdb')
                cur.execute(create_db)
                cur.execute(executelock)
                cur.execute(statement)
                while cur.nextset():
                    pass
                cur.execute(releaselock)
                db_connection = pyodbc.connect(  # need use back the recovered db
                    'DRIVER={ODBC Driver 17 for SQL Server}; \
                    SERVER=' + 'DESKTOP-75MSPGF\MSSQL15.MSSQLSERVER01' + '; \
                    DATABASE=' + 'database1' + ';\
                    Trusted_Connection=yes;' \
                    , autocommit=True)
                print("Update and restore secondary db ")
                print("Using second db!")
        except:  # if secondary db was deleted, it will auto restore
            db_connection = pyodbc.connect(  #
                'DRIVER={ODBC Driver 17 for SQL Server}; \
                SERVER=' + 'DESKTOP-75MSPGF\MSSQL15.MSSQLSERVER01' + '; \
                Trusted_Connection=yes;' \
                , autocommit=True)
            cur = db_connection.cursor()

            executelock = ("alter database database1 set offline with rollback immediate ")
            releaselock = ("alter database database1 set online")
            folder_path = "C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\"
            folders = os.listdir(folder_path)  # ['folder1', 'folder2'] list folder
            statement = (
                f"RESTORE DATABASE database1 FROM  DISK = 'C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\{folders[-1]}' WITH RECOVERY, MOVE 'database' TO 'C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\MSSQL15.SQLEXPRESS01\\MSSQL\DATA\\database.mdf',MOVE 'database_log' TO 'C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\MSSQL15.SQLEXPRESS01\\MSSQL\DATA\\database_log.ldf', REPLACE")  # file name and db must be the same
            create_db = (
                "IF NOT EXISTS( SELECT * FROM sys.databases WHERE name = 'database1' ) BEGIN CREATE DATABASE database1 ; END;")
            print(cur.execute(create_db), 'createdb')
            cur.execute(create_db)
            cur.execute(executelock)
            cur.execute(statement)
            while cur.nextset():
                pass
            cur.execute(releaselock)
            db_connection = pyodbc.connect(  # need use back the recovered db
                'DRIVER={ODBC Driver 17 for SQL Server}; \
                SERVER=' + 'DESKTOP-75MSPGF\SMSSQL15.MSSQLSERVER01' + '; \
                DATABASE=' + 'database1' + ';\
                Trusted_Connection=yes;' \
                , autocommit=True)
            print("Restored second db as it was never created.")
            print("Using second db!")
            print(';')

    # return cur
    return db_connection


def doctor_and_patient_needed(needpatientandresearcher):
    @wraps(needpatientandresearcher)
    def decorated_func(*args, **kwargs):
        print(session)
        if "access_level" in session:
            if session['access_level'] in ["doctor","patient"]:
                print("Hello")
                return needpatientandresearcher(*args, **kwargs)
        flash("Invalid access","error")
        return redirect(url_for('homepage'))
    return decorated_func


def researcher_needed(needresearcher):
    @wraps(needresearcher)
    def decorated_func(*args, **kwargs):
        if session['access'] != '0':
            flash("Invalid User")
            return redirect(url_for('home'))
        else:
            return needresearcher(*args, **kwargs)

    return decorated_func


def patient_needed(needpatient):
    @wraps(needpatient)
    def decorated_func(*args, **kwargs):
        if session['access'] != '1':
            flash("Invalid User")
            return redirect(url_for('home'))
        else:
            return needpatient(*args, **kwargs)

    return decorated_func


def doctor_needed(needdoctor):
    @wraps(needdoctor)
    def decorated_func(*args, **kwargs):
        if session['access_level'] != 'doctor':
            flash("Invalid User","error")
            return redirect(url_for('homepage'))
        else:
            return needdoctor(*args, **kwargs)
    return decorated_func


def hr_needed(needHR):
    @wraps(needHR)
    def decorated_func(*args, **kwargs):
        if session['access_level'] != 'hr':
            flash("Invalid User","error")
            return redirect(url_for('homepage'))
        else:
            return needHR(*args, **kwargs)
    return decorated_func


def admin_needed(needadmin):
    @wraps(needadmin)
    def decorated_func(*args, **kwargs):
        if session['access'] < '3':
            flash("Invalid User")
            return redirect(url_for('home'))
        else:
            return needadmin(*args, **kwargs)

    return decorated_func


def head_admin_needed(needhadmin):
    @wraps(needhadmin)
    def decorated_func(*args, **kwargs):
        if session['access'] < '4':
            flash("Invalid User")
            return redirect(url_for('home'))
        else:
            return needhadmin(*args, **kwargs)

    return decorated_func

def allowed_filename(filename):
    expression=re.compile(r"(?i)^[\w]*(.docx)$")
    return re.fullmatch(expression,filename)

def check_file_hash(file_name,stored_hash):
    with open(os.path.join(app.config['UPLOAD_FOLDER'], file_name),"rb") as original_file:
        md5Hash = hashlib.md5(original_file.read())
        fileHashed = md5Hash.hexdigest()
        return fileHashed==stored_hash

def get_file_data_from_database(patient_id):
    connection = auto_use_seconddb()

    cursor = connection.cursor()
    data = cursor.execute("select * from patient_file where patient_id=?",(patient_id)).fetchone()
    cursor.close()
    return data



@app.context_processor
def inject_templates_with_session_date():
    try:
        print('no session[\'db\'')
        return dict(session, server = session['db'])
    except:
        return dict(session)
        

with app.app_context():
    @app.route('/homepage')
    @custom_login_required
    def homepage():
        flash("welcome")
        cnxn = auto_use_seconddb()
        if session['access_level'] == 'patient' or session['access_level'] == 'doctor' or session['access_level'] == 'researcher':
            cursor = cnxn.cursor()
            cursor.execute("select file_content from patient_file where patient_id = ?", (session['id']))
            data = cursor.fetchall()
            if data == None:
                return render_template(url_for('baseinfo'))
            return render_template('homepage.html')
        elif session['access_level'] == 'head_admin':
            return redirect(url_for('dashboard'))
        elif session['access_level'] == 'hr':
            return redirect(url_for('dashboard'))

    @app.route('/')
    def index():
        session.clear()
        return render_template('index.html')

    @app.route('/dashboard')
    @custom_login_required
    def dashboard():
        if session['access_level'] == 'head_admin' or session['access_level'] == 'hr':
            cnxn = auto_use_seconddb()
            cursor = cnxn.cursor()
            patients = cursor.execute("SELECT * FROM patients").fetchall()
            doctors = cursor.execute("SELECT * FROM doctors").fetchall()
            hr = cursor.execute("SELECT * FROM hr").fetchall()
            researcher = cursor.execute("SELECT * FROM researchers").fetchall()
            return render_template('dashboard.html', patients=patients, doctors=doctors, hr=hr, researchers=researcher)
        else:
            try:
                filter = ac_log('dashboard',request.remote_addr,session['username'])
                ac_logLogger.addFilter(filter)
                ac_logLogger.warning(
                    msg=f"code 401"
                )
            except:
                filter = ac_log('dashboard',request.remote_addr,session['username'])
                ac_logLogger.addFilter(filter)
                ac_logLogger.warning(
                    msg=f"code 401"
                )
            return redirect(url_for('access_denied'))

    @app.route('/401')
    def access_denied():
        if session['login'] != True:
            return render_template('401.html')
        else:
            flash('Access denied','error')
            return redirect(url_for('homepage'))


    @app.route('/404')
    def page_not_found():
        return render_template('404.html')


    @app.route('/500')
    def server_error():
        return render_template('500.html')


    @app.route('/table')
    @custom_login_required
    def table():
        return render_template('table.html')

    def SendMail(ImgFileName, email,topic):
        with open(ImgFileName, 'rb') as f:
            img_data = f.read()
        sender = "IT2566proj@gmail.com"
        senderpass = 'FishNugget123'
        msg = MIMEMultipart()
        msg['Subject'] = topic+' Qr Code For Google Authenticator'
        msg['From'] = 'AngelHealth@mail.gov.sg'
        msg['To'] = email

        text = MIMEText("test")
        msg.attach(text)
        image = MIMEImage(img_data, name=os.path.basename(ImgFileName))
        msg.attach(image)

        s = smtplib.SMTP('smtp.gmail.com:587')
        s.ehlo()
        s.starttls()
        s.ehlo()
        s.login(sender, senderpass)
        s.sendmail(sender, email, msg.as_string())
        s.quit()


    def add_admin():
        while True:
            key = input("Do you want to create Head Admin ID and password? (Y/N/Show/Delete)").capitalize()
            if key == "Y":
                connection = auto_use_seconddb()
                
                cursor = connection.cursor()
                pattern = ('^\d{6}[A-Za-z]$')
                username = input("Enter New Head Admin ID: ")
                result = re.match(pattern,username)
                while result == None:
                    print("Only first 6 digits and 1 alphabet at the end!")
                    username = input("Enter New Head Admin ID: ")
                    result = re.match(pattern,username)


                check = cursor.execute("SELECT username FROM head_admin WHERE username = ?",
                                       (username)).fetchval()  # prevent sql injection

                firstname = input("Enter New Head Admin First Name: ")
                lastname = input("Enter New Head Admin last Name: ")
                email = input("Enter New Head Admin email: ")

                pattern = ('^(?=\S{10,20}$)(?=.*?\d)(?=.*?[a-z])(?=.*?[A-Z])(?=.*?[^A-Za-z\s0-9])')
                admin_password = input("Enter New Head Admin Password: ")

                result = re.match(pattern,admin_password)
                while result == None:
                    print( "Password must contain 10-20 characters, number, uppercase, lowercase, special character.")
                    admin_password = input("Enter New Head Admin Password: ")
                    result = re.match(pattern,admin_password)
                phone_no = input("Enter New Head Admin Phone Number: ")
                
                otp_code = pyotp.random_base32()

                md5Hash = hashlib.md5(admin_password.encode("utf-8"))
                md5Hashed = md5Hash.hexdigest()
                address = input("Enter New Head Admin address: ")
                postal_code = input("Enter New Head Admin postal code: ")

                cursor = connection.cursor()
                check_username = cursor.execute("SELECT username FROM head_admin WHERE username = ?",
                                       (username)).fetchval()  # prevent sql injection
                check_email = cursor.execute("SELECT email FROM head_admin WHERE email = ?",
                                       (email)).fetchval()  # prevent sql injection

                if check_email == None and check_username == None :
                    insert_query = "INSERT INTO head_admin (username, first_name, last_name, pass_hash,email,otp_code,phone_no,access_level,postal_code,address) \
                            VALUES (?, ?, ?, ?, ?, ?,?,?,?,?); "
                    values = (username, firstname, lastname, md5Hashed,
                              email,otp_code,phone_no,'head_admin',postal_code,address)
                    cursor.execute(insert_query, values)
                    cursor.commit()
                    # connection = auto_use_seconddb()
                    # cursor = connection.cursor()
                    insert_query = "INSERT INTO access_list (username,access_level,pass_hash) \
                            VALUES (?, ?,?); "
                    values = (username,'head_admin',md5Hashed)
                    cursor.execute(insert_query, values)
                    cursor.commit()
                    cursor.close()
                    print('Sending email OTP...')
                    qr = 'otpauth://totp/AngelHealth:' + str(username) + '?secret=' + otp_code
                    image = pyqrcode.create(qr)
                    image.png('image.png', scale=5)
                    SendMail("image.png", email, "Head Admin")
                    os.remove('image.png')
                    print('Successful creating Head Admin')
                    filter = lr_log(username)
                    lr_logLogger.addFilter(filter)
                    lr_logLogger.debug(
                        msg=f"successfully registered as Head Admin staff with Angel Health"
                    )
                    cursor = connection.cursor()
                    check = cursor.execute("SELECT * FROM head_admin").fetchall()  # prevent sql injection
                    print("List of Head Admins:")
                    for x in check:
                        username = x.username.strip()
                        first_name = x.first_name.strip()
                        last_name = x.last_name.strip()
                        email = x.email.strip()
                        phone_no = x.phone_no.strip()
                        print(f"Username: {username}, First Name: {first_name}, Last Name: {last_name}, Email: {email}, Phone No: {phone_no}")
                    continue
                else:
                    print("Head Admin already exists! Check your username and email!!!!")
      
            elif key == "N":
                break
            elif key == "Show":
                connection = auto_use_seconddb()
                cursor = connection.cursor()
                check = cursor.execute("SELECT * FROM head_admin").fetchall()  # prevent sql injection
                print("List of Head Admins:")
                for x in check:
                    username = x.username.strip()
                    first_name = x.first_name.strip()
                    last_name = x.last_name.strip()
                    email = x.email.strip()
                    print(f"Username: {username}, First Name: {first_name}, Last Name: {last_name}, Email: {email}")

                cursor.close()
            elif key == 'Delete':
                # try:
                connection = auto_use_seconddb()
                cursor = connection.cursor()
                key = input("Enter the Head Admin ID to delete: ")
                check = cursor.execute("SELECT * FROM head_admin WHERE username = ?",
                                       (key)).fetchval()  # prevent sql injection
                if check == None:
                    print("Head admin does not exist")
                    continue

                else:
                    check = cursor.execute("DELETE FROM head_admin WHERE username = ?", (key))  # prevent sql injection
                    cursor.commit()
                    check = cursor.execute("DELETE FROM access_list WHERE username = ?", (key))  # prevent sql injection
                    cursor.commit()
                    print(f"{key} was removed as Head Admin.")
                    filter = lr_log(key)
                    lr_logLogger.addFilter(filter)
                    lr_logLogger.debug(
                        msg=f",Head Admin staff has been removed from Angel Health"
                    )
                    check = cursor.execute("SELECT * FROM head_admin").fetchall()  # prevent sql injection
                    print("List of Head Admins:")
                    for x in check:
                        username = x.username.strip()
                        first_name = x.first_name.strip()
                        last_name = x.last_name.strip()
                        email = x.email.strip()
                        print(f"Username: {username}, First Name: {first_name}, Last Name: {last_name}, Email: {email}")
                    cursor.close()


            # except:
            #     print('Error in deleting Head Admin in MSSQL Database')
            else:
                print("Please enter Y or N or Delete only!")
                continue

    # must add get method to to retrieve url
    @app.route("/viewUser", methods=["GET", "POST"])
    def viewUser():
        if request.method == 'GET':
            print(session)
            return render_template("customerDetail.html")

    @app.route('/updateUser', methods=['GET', 'POST'])
    def update_user():
        update_user_form = General_UpdateForm(request.form)
        if request.method == 'POST' and update_user_form.validate() and session['access_level'] == 'patient':  # POST method, update upon clicking submission
            cnxn = auto_use_seconddb()
            
            cursor = cnxn.cursor()
            firstname = update_user_form.firstname.data
            lastname = update_user_form.lastname.data
            address = update_user_form.address.data
            phone_no = update_user_form.phone_no.data
            postal_code = update_user_form.postal_code.data
            email = update_user_form.email.data

            insert_query1 = ('UPDATE patients SET first_name = ?,last_name = ?,address = ?,phone_no = ?,postal_code = ?,email = ? WHERE username = ?;')
            
            values1 = (firstname, lastname, address,phone_no,postal_code,email,session['username'])


        

            if update_user_form.password.data == '' and update_user_form.confirmPassword.data == '':
                print('did not update password')
                pass
            else:
                md5Hash = hashlib.md5(update_user_form.password.data.encode("utf-8"))
                md5Hashed = md5Hash.hexdigest()
                insert_query2 = ('UPDATE patients  SET pass_hash = ? WHERE username = ?;')
            
                values2 = (md5Hashed,session['username'])
              
                insert_query3 = textwrap.dedent('''
                    UPDATE access_list SET pass_hash = ? WHERE username = ?; 
                ''')
                values3 = (md5Hashed,session['username'])
                cursor.execute(insert_query2, values2)
                cursor.execute(insert_query3, values3)
            cursor.execute(insert_query1, values1)
            print(cursor.execute(insert_query1, values1),'check up date not pw')
            
            cursor.commit()

            print('updated')
            cursor = cnxn.cursor()

        
            user_info = cursor.execute(
            "select * from patients where username = ?",(session['username'])).fetchone() #fetchone() dont delete this except others
            session['id'] = user_info[0]
            session['username'] = user_info[1].strip()
            session['first_name'] = user_info[2].strip()
            session['last_name'] = user_info[3].strip()
            session['phone_no'] =  user_info[7].strip()
            session['email'] =  user_info[6].strip()
            session['address'] =  user_info[8].strip()
            session['postal_code'] =  user_info[9].strip()
            session['access_level'] = user_info[13].strip()

            print(user_info)
            print(session['id'],'this is session')
            filter = lr_log(session['username'])
            lr_logLogger.addFilter(filter)
            lr_logLogger.debug(
                msg=f"has updated his/her profile"
            )
            cursor.close()

            return redirect(url_for('viewUser'))


        # to retrieve data from shelve and diplay previous data
        # so bascially this will come first as admin did not click update thus post doesnt work first
        elif request.method == 'POST' and update_user_form.validate() and session['access_level'] == 'head_admin': 
            cnxn = auto_use_seconddb()
            
            cursor = cnxn.cursor()
            firstname = update_user_form.firstname.data
            lastname = update_user_form.lastname.data
            address = update_user_form.address.data
            phone_no = update_user_form.phone_no.data
            postal_code = update_user_form.postal_code.data
            email = update_user_form.email.data

            insert_query1 = ('UPDATE head_admin SET first_name = ?,last_name = ?,address = ?,phone_no = ?,postal_code = ?,email = ? WHERE username = ?;')
            
            values1 = (firstname, lastname, address,phone_no,postal_code,email,session['username'])

            if update_user_form.password.data == '' and update_user_form.confirmPassword.data == '':
                print('did not update [asswprd')
                pass
            else:
                md5Hash = hashlib.md5(update_user_form.password.data.encode("utf-8"))
                md5Hashed = md5Hash.hexdigest()
                insert_query2 = ('UPDATE head_admin SET pass_hash = ? WHERE username = ?;')
            
                values2 = (md5Hashed,session['username'])
              
                insert_query3 = textwrap.dedent('''
                    UPDATE access_list 
                    SET pass_hash = ? \
                    WHERE username = ?; 
                ''')
                values3 = (md5Hashed,session['username'])
                cursor.execute(insert_query2, values2)
                cursor.execute(insert_query3, values3)
            cursor.execute(insert_query1, values1)
            
            cursor.commit()

            print('updated')
            cursor = cnxn.cursor()

            user_info = cursor.execute(
            "select * from head_admin where username = ?",(session['username'])).fetchone() #fetchone() dont delete this except others
            session['id'] = user_info[0]
            session['username'] = user_info[1].strip()
            session['first_name'] = user_info[2].strip()
            session['last_name'] = user_info[3].strip()
            session['phone_no'] =  user_info[7].strip()
            session['email'] =  user_info[6].strip()
            session['address'] =  user_info[10].strip()
            session['postal_code'] =  user_info[9].strip()
            session['access_level'] =  user_info[8].strip()
            print(user_info)
            print(session['id'],'this is session')
            filter = lr_log(session['username'])
            lr_logLogger.addFilter(filter)
            lr_logLogger.debug(
                msg=f"has updated his/her profile"
            )
            cursor.close()

            return redirect(url_for('viewUser'))
        elif request.method == 'POST' and update_user_form.validate() and session['access_level'] == 'hr': 
            cnxn = auto_use_seconddb()
            
            cursor = cnxn.cursor()
            firstname = update_user_form.firstname.data
            lastname = update_user_form.lastname.data
            address = update_user_form.address.data
            phone_no = update_user_form.phone_no.data
            postal_code = update_user_form.postal_code.data
            email = update_user_form.email.data
            insert_query1 = ('UPDATE hr  SET first_name = ?, last_name = ?,address = ?,phone_no = ?,postal_code = ?,email = ? WHERE username = ?;')
            # insert_query1 = ('UPDATE patients SET first_name = ?,last_name = ?,address = ?,phone_no = ?,postal_code = ?,email = ? WHERE username = ?;')
            
            values1 = (firstname, lastname, address,phone_no,postal_code,email,session['username'])


        

            if update_user_form.password.data == '' and update_user_form.confirmPassword.data == '':
                print('did not update [asswprd')
                pass
            else:
                md5Hash = hashlib.md5(update_user_form.password.data.encode("utf-8"))
                md5Hashed = md5Hash.hexdigest()
                insert_query2 = (' UPDATE hr SET pass_hash = ? WHERE username = ?;')
            
                values2 = (md5Hashed,session['username'])
              
                insert_query3 = textwrap.dedent('''
                    UPDATE access_list SET pass_hash = ? WHERE username = ?; 
                ''')
                values3 = (md5Hashed,session['username'])
                cursor.execute(insert_query2, values2)
                cursor.execute(insert_query3, values3)
            cursor.execute(insert_query1, values1)
            print(cursor.execute(insert_query1, values1),'check up date not pw')
            
            cursor.commit()

            print('updated')
            cursor = cnxn.cursor()

            user_info = cursor.execute(
            "select * from hr where username = ?",(session['username'])).fetchone() #fetchone() dont delete this except others
            session['id'] = user_info[0]
            session['username'] = user_info[1].strip()
            session['first_name'] = user_info[2].strip()
            session['last_name'] = user_info[3].strip()
            session['phone_no'] =  user_info[9].strip()
            session['email'] =  user_info[8].strip()
            session['address'] =  user_info[6].strip()
            session['postal_code'] =  user_info[7].strip()
            session['access_level'] =  user_info[10].strip()
            print(user_info)
            print(session['id'],'this is session')

            cursor.close()
            filter = lr_log(session['username'])
            lr_logLogger.addFilter(filter)
            lr_logLogger.debug(
                msg=f"has updated his/her profile"
            )
            return redirect(url_for('viewUser'))
        elif request.method == 'POST' and update_user_form.validate() and session['access_level'] == 'doctor': 
            cnxn = auto_use_seconddb()
            
            cursor = cnxn.cursor()
            firstname = update_user_form.firstname.data
            lastname = update_user_form.lastname.data
            address = update_user_form.address.data
            phone_no = update_user_form.phone_no.data
            postal_code = update_user_form.postal_code.data
            email = update_user_form.email.data
            insert_query1 = ('UPDATE doctors SET first_name = ?,last_name = ?,address = ?,phone_no = ?,postal_code = ?,email = ? WHERE username = ?;')
            # insert_query1 = ('UPDATE patients SET first_name = ?,last_name = ?,address = ?,phone_no = ?,postal_code = ?,email = ? WHERE username = ?;')
            
            values1 = (firstname, lastname, address,phone_no,postal_code,email,session['username'])


        

            if update_user_form.password.data == '' and update_user_form.confirmPassword.data == '':
                print('did not update [asswprd')
                pass
            else:
                md5Hash = hashlib.md5(update_user_form.password.data.encode("utf-8"))
                md5Hashed = md5Hash.hexdigest()
                insert_query2 = (' UPDATE doctors SET pass_hash = ? WHERE username = ?;')
            
                values2 = (md5Hashed,session['username'])
              
                insert_query3 = textwrap.dedent('''
                    UPDATE access_list SET pass_hash = ? WHERE username = ?; 
                ''')
                values3 = (md5Hashed,session['username'])
                cursor.execute(insert_query2, values2)
                cursor.execute(insert_query3, values3)
            cursor.execute(insert_query1, values1)
            print(cursor.execute(insert_query1, values1),'check up date not pw')
            
            cursor.commit()

            print('updated')
            cursor = cnxn.cursor()

            user_info = cursor.execute(
            "select * from doctors where username = ?",(session['username'])).fetchone() #fetchone() dont delete this except others
            session['id'] = user_info[0]
            session['username'] = user_info[1].strip()
            session['first_name'] = user_info[2].strip()
            session['last_name'] = user_info[3].strip()
            session['phone_no'] =  user_info[11].strip()
            session['email'] =  user_info[6].strip()
            session['address'] =  user_info[7].strip()
            session['postal_code'] =  user_info[8].strip()
            session['access_level'] =  user_info[10].strip()
            print(user_info)
            print(session['id'],'this is session')
            filter = lr_log(session['username'])
            lr_logLogger.addFilter(filter)
            lr_logLogger.debug(
                msg=f"has updated his/her profile"
            )
            cursor.close()

            return redirect(url_for('viewUser'))
        elif request.method == 'POST' and update_user_form.validate() and session['access_level'] == 'researcher': 
            cnxn = auto_use_seconddb()
            cursor = cnxn.cursor()
            firstname = update_user_form.firstname.data
            lastname = update_user_form.lastname.data
            address = update_user_form.address.data
            phone_no = update_user_form.phone_no.data
            postal_code = update_user_form.postal_code.data
            email = update_user_form.email.data
            insert_query1 = ('UPDATE researchers SET first_name = ?,last_name = ?,address = ?,phone_no = ?,postal_code = ?,email = ? WHERE username = ?;')
            # insert_query1 = ('UPDATE patients SET first_name = ?,last_name = ?,address = ?,phone_no = ?,postal_code = ?,email = ? WHERE username = ?;')
            
            values1 = (firstname, lastname, address,phone_no,postal_code,email,session['username'])


        

            if update_user_form.password.data == '' and update_user_form.confirmPassword.data == '':
                print('did not update [asswprd')
                pass
            else:
                md5Hash = hashlib.md5(update_user_form.password.data.encode("utf-8"))
                md5Hashed = md5Hash.hexdigest()
                insert_query2 = ('UPDATE researchers SET pass_hash = ? WHERE username = ?;')
            
                values2 = (md5Hashed,session['username'])
              
                insert_query3 = textwrap.dedent('''
                    UPDATE access_list SET pass_hash = ? WHERE username = ?; 
                ''')
                values3 = (md5Hashed,session['username'])
                cursor.execute(insert_query2, values2)
                cursor.execute(insert_query3, values3)
            cursor.execute(insert_query1, values1)
            print(cursor.execute(insert_query1, values1),'check up date not pw')
            
            cursor.commit()

            print('updated')
            cursor = cnxn.cursor()

            user_info = cursor.execute(
            "select * from researchers where username = ?",(session['username'])).fetchone() #fetchone() dont delete this except others
            session['id'] = user_info[10]
            session['username'] = user_info[0].strip()
            session['first_name'] = user_info[1].strip()
            session['last_name'] = user_info[2].strip()
            session['phone_no'] =  user_info[6].strip()
            session['email'] =  user_info[5].strip()
            session['address'] =  user_info[7].strip()
            session['postal_code'] =  user_info[8].strip()
            session['access_level'] =  user_info[11].strip()
            print(user_info)
            print(session['id'],'this is session')
            filter = lr_log(session['username'])
            lr_logLogger.addFilter(filter)
            lr_logLogger.debug(
                msg=f"has updated his/her profile"
            )
            cursor.close()

            return redirect(url_for('viewUser'))


        elif request.method == 'GET':  # get method
            cnxn = auto_use_seconddb()
            if session['access_level'] == 'patient':
                cursor = cnxn.cursor()
                user_info = cursor.execute(
                    "select * from patients where username = ?",
                    (session['username'])).fetchone() #fetchone() dont delete this except others
           
                get_firstname = user_info[2].strip()
                get_lastname = user_info[3].strip()
                get_phone_no =  user_info[7]
                get_email =  user_info[6].strip()
                get_address =  user_info[8].strip()
                get_postal_code =  user_info[9]
                

            elif session['access_level'] == 'researcher':
                cursor = cnxn.cursor()
                user_info = cursor.execute(
                    "select * from researchers where username = ?",
                    (session['username'])).fetchone() #fetchone() dont delete this except others
                get_firstname = user_info[1].strip()
                get_lastname = user_info[2].strip()
                get_phone_no =  user_info[6].strip()
                get_email =  user_info[5].strip()
                get_address =  user_info[7].strip()
                get_postal_code =  user_info[8].strip()


            elif session['access_level'] == 'doctor':
                cursor = cnxn.cursor()
                user_info = cursor.execute(
                    "select * from doctors where username = ?",
                    (session['username'])).fetchone() #fetchone() dont delete this except others
                get_firstname = user_info[2].strip()
                get_lastname = user_info[3].strip()
                get_phone_no =  user_info[11].strip()
                get_email =   user_info[6].strip()
                get_address =  user_info[7].strip()
                get_postal_code =  user_info[8].strip()

            elif session['access_level'] == 'hr':
                cursor = cnxn.cursor()
                user_info = cursor.execute(
                    "select * from hr where username = ?",
                    (session['username'])).fetchone() #fetchone() dont delete this except others
                get_firstname = user_info[2].strip()
                get_lastname = user_info[3].strip()
                get_phone_no =  user_info[9].strip()
                get_email =  user_info[8].strip()
                get_address =  user_info[6].strip()
                get_postal_code =  user_info[7].strip()

            elif session['access_level'] == 'head_admin':
                cursor = cnxn.cursor()
                user_info = cursor.execute(
                    "select * from head_admin where username = ?",
                    (session['username'])).fetchone() #fetchone() dont delete this except others
                get_firstname = user_info[2].strip()
                get_lastname = user_info[3].strip()
                get_phone_no =  user_info[7].strip()
                get_email =  user_info[6].strip()
                get_address =  user_info[10].strip()
                get_postal_code =  user_info[9].strip()
            print(user_info,'user_id')
            print(cursor,'cursor')


            update_user_form.firstname.data =get_firstname
            update_user_form.lastname.data = get_lastname
            update_user_form.address.data = get_address
            update_user_form.phone_no.data = get_phone_no
            update_user_form.postal_code.data = get_postal_code
            update_user_form.email.data = get_email


            return render_template("updateUser.html", form = update_user_form)
        return render_template("updateUser.html", form = update_user_form)

    # @auto_use_seconddb
    @app.route('/login', methods=['GET', 'POST'])
    def login():
        patient_login_form = Patient_Login_form(request.form)
        admin_login_form = Admin_Login_form(request.form)
        cursor = auto_use_seconddb().cursor()
        if patient_login_form.patient_submit.data and patient_login_form.validate():
            passed = False
            username = patient_login_form.username.data
            password = patient_login_form.password.data
            print(username,password)
            md5Hash = hashlib.md5(password.encode("utf-8"))
            md5Hashed = md5Hash.hexdigest()
            print(md5Hashed)
            user_info = cursor.execute(
                "select * from patients where username = ? and pass_hash = ?",
                (username, md5Hashed)).fetchone() #fetchone() dont delete this except others
            print(user_info,'user_id')
            print(cursor,'cursor')
            # for reference list index!
            # /****** Script for SelectTopNRows command from SSMS  ******/
            # SELECT TOP (1000) [patient_id] 0
            #     ,[username] 1
            #     ,[first_name]2
            #     ,[last_name]3
            #     ,[pass_hash]4
            #     ,[otp_code]5
            #     ,[email]6
            #     ,[phone_no]7
            #     ,[address]8
            #     ,[postal_code]9
            #     ,[hospital]10
            #     ,[tending_physician]11
            #     ,[appointment]12
            #     ,[access_level]13
            # FROM [database1].[dbo].[patients]
            if user_info != None:

                passed = True
                print(user_info[0],'id')
                print(user_info[1].strip(),'id')
                print(user_info[3].strip(),'id')
                print(user_info[4].strip(),'id')
                print(user_info[7],'id')
                session['id'] = user_info[0]
                session['username'] = user_info[1].strip()
                session['first_name'] = user_info[2].strip()
                session['last_name'] = user_info[3].strip()
                session['phone_no'] =  user_info[7]
                session['email'] =  user_info[6].strip()
                session['address'] =  user_info[8].strip()
                session['postal_code'] =  user_info[9]
                session['access_level'] = user_info[13].strip()

                # should we add this too? as session?
                try:# if it is not None then  user_info[3]
                    session['tending_physician'] = user_info[11].strip()
                except:
                    session['tending_physician'] = None
                try:# if it is tending_physician None then strip
                    session['appointment'] = user_info[12].strip()
                except:
                    session['appointment'] =  None
                
                session['otp-semi-login'] = True
                print('user logged in')
                flash("Please continue with otp validation before login is successful","success")
                filter = lr_log(session['username'])
                lr_logLogger.addFilter(filter)
                lr_logLogger.debug(
                    msg=f"has been redirected to otpvalidation page"
                )

            else:
                print("ERROR")
                flash("Invalid username or password","error")
                filter = lr_log(patient_login_form.username.data)
                lr_logLogger.addFilter(filter)
                lr_logLogger.debug(
                    msg=f"has failed to be redirected to otpvalidation page"
                )
                return redirect(url_for('login'))

        elif admin_login_form.staff_submit.data and admin_login_form.validate():
            passed = False
            username = admin_login_form.username.data
            password = admin_login_form.password.data
            md5Hash = hashlib.md5(password.encode("utf-8"))
            md5Hashed = md5Hash.hexdigest()
            # cursor = auto_use_seconddb().cursor()
            print(cursor)
            access_info = cursor.execute(
                "select * from access_list where username = ? and pass_hash = ?",
                (username, md5Hashed)).fetchone() #fetchone() dont delete this except others


            #added md5Hashed password in access_list table so to user username and hashed password for checking if it is in both access_list table and unqiue role tables
            #check whether in access list table
            if access_info != None:

                # cnxn = auto_use_seconddb()
            
                # session['username'] = access_info[0].strip()
                # session['access_list'] = access_info[1].strip()
                access_level_username = access_info[0].strip()
                access_level_access_list = access_info[1].strip()
                print(access_level_username,'username')
                print(access_level_access_list)
                #check whether was it created in the individual tables (e.g. is it in doctors or researchers or HR or head_admin tables?)
                #Retrieve the unique roles using access_level_access_list
                if access_level_access_list == "researcher":
                    user_info = cursor.execute(
                    f"select * from researchers where username = ? and pass_hash = ?",
                    (access_level_username, md5Hashed)).fetchone() #fetchone() dont delete this except others
                    # CREATE TABLE [dbo].[researchers](
                    #     [username] [nchar](30) NOT NULL,0
                    #     [first_name] [nchar](30) NOT NULL,1
                    #     [last_name] [nchar](30) NOT NULL,2
                    #     [pass_hash] [varchar](50) NOT NULL,3
                    #     [otp_code] [varchar](50) NOT NULL,4
                    #     [email] [varchar](50) NOT NULL,5
                    #     [phone_no] [varchar](20) NOT NULL,6
                    #     [address] [varchar](50) NOT NULL,7
                    #     [postal_code] [varchar](6) NOT NULL,8
                    #     [company] [varchar](50) NULL,9
                    #     [researcher_id] [int] IDENTITY(1,1) NOT NULL,10
                    #     [access_level][varchar](30) NOT NULL 11
                    # ) ON [PRIMARY]
                    if user_info != None:
                        passed = True
                        print(user_info[10],'researcher_id')
                        print(user_info[0].strip(),'username')
                        print(user_info[2].strip(),'last_name')
                        print(user_info[6].strip(),'phone_no')
                        print(user_info[7].strip(),'address')
                        print(user_info[8].strip(),'postal_code')
                        print(user_info[11].strip(),'access_level')

                        session['id'] = user_info[10]
                        session['username'] = user_info[0].strip()
                        session['first_name'] = user_info[1].strip()
                        session['last_name'] = user_info[2].strip()
                        session['phone_no'] =  user_info[6].strip()
                        session['email'] =  user_info[5].strip()
                        session['address'] =  user_info[7].strip()
                        session['postal_code'] =  user_info[8].strip()
                        session['access_level'] =  user_info[11].strip()

                        # should we add this too? as session?
                        try:
                            session['company'] = user_info[9].strip() # If HR dk the company
                        except:
                            session['company'] = None
                elif access_level_access_list == "head_admin":
                    user_info = cursor.execute(
                    f"select * from {access_level_access_list} where username = ? and pass_hash = ?",
                    (access_level_username, md5Hashed)).fetchone() #fetchone() dont delete this except others
                    # cursor = cnxn.close()
                        # CREATE TABLE [dbo].[head_admin](
                        # 	[head_admin_id] [int] IDENTITY(1,1) NOT NULL, 0
                        # 	[username] [nchar](30) NOT NULL,1
                        # 	[first_name] [nchar](30) NOT NULL,2
                        # 	[last_name] [nchar](30) NOT NULL,3
                        # 	[pass_hash] [varchar](50) NOT NULL,4
                        # 	[otp_code] [varchar](50) NOT NULL,5
                        # 	[email] [varchar](50) NOT NULL,6
                        # 	[phone_no] [varchar](20) NOT NULL,7
                        # 	[access_level][varchar](30) NOT NULL8
                        #   [postal_code][varchar](6) NOT NULL,9
	                    #   [address][varchar](50) NOT NULL10
                        # ) ON [PRIMARY]
                    if user_info != None:
                        passed = True
                        print(user_info[0],'id')
                        print(user_info[1].strip(),'username')
                        print(user_info[2].strip(),'first_name')
                        print(user_info[3].strip(),'last_name')
                        print(user_info[7].strip(),'phone_no')
                        print(user_info[6].strip(),'email')
                        print(user_info[10].strip(),'address')
                        print(user_info[9].strip(),'postal_code')
                        print(user_info[8].strip(),'access_level')

                        session['id'] = user_info[0]
                        session['username'] = user_info[1].strip()
                        session['first_name'] = user_info[2].strip()
                        session['last_name'] = user_info[3].strip()
                        session['phone_no'] =  user_info[7].strip()
                        session['email'] =  user_info[6].strip()
                        session['address'] =  user_info[10].strip()
                        session['postal_code'] =  user_info[9].strip()
                        session['access_level'] =  user_info[8].strip()
                        

                elif access_level_access_list == "hr":
                    user_info = cursor.execute(
                    f"select * from hr where username = ? and pass_hash = ?",
                    (access_level_username, md5Hashed)).fetchone() #fetchone() dont delete this except others
                    # cursor = cnxn.close()
                    # CREATE TABLE [dbo].[hr](
                    #     [hr_id] [int] IDENTITY(1,1) NOT NULL,0
                    #     [username] [nchar](10) NOT NULL,1
                    #     [first_name] [nchar](20) NOT NULL,2
                    #     [last_name] [nchar](20) NOT NULL,3
                    #     [pass_hash] [varchar](50) NOT NULL,4
                    #     [otp_code] [varchar](50) NOT NULL,5
                    #     [address] [varchar](50) NOT NULL,6
                    #     [postal_code] [varchar](6) NOT NULL,7
                    #     [email] [varchar](50) NOT NULL,8
                    #     [phone_no] [varchar](20) NOT NULL,9
                    #     [access_level][varchar](30) NOT NULL 10
                    # ) ON [PRIMARY]
                    if user_info != None:
                        passed = True
                        print(user_info[0],'id')
                        print(user_info[1].strip(),'username')
                        print(user_info[2].strip(),'first_name')
                        print(user_info[3].strip(),'last_name')
                        print(user_info[9].strip(),'phone_no')
                        print(user_info[8].strip(),'email')
                        print(user_info[6].strip(),'address')
                        print(user_info[7].strip(),'postal_code')
                        print(user_info[10].strip(),'access_level')

                        session['id'] = user_info[0]
                        session['username'] = user_info[1].strip()
                        session['first_name'] = user_info[2].strip()
                        session['last_name'] = user_info[3].strip()
                        session['phone_no'] =  user_info[9].strip()
                        session['email'] =  user_info[8].strip()
                        session['address'] =  user_info[6].strip()
                        session['postal_code'] =  user_info[7].strip()
                        session['access_level'] =  user_info[10].strip()

                elif access_level_access_list == "doctor":
                    user_info = cursor.execute(
                    f"select * from doctors where username = ? and pass_hash = ?",
                    (access_level_username, md5Hashed)).fetchone() #fetchone() dont delete this except others
                    # cursor = cnxn.close()
                    # CREATE TABLE [dbo].[doctors](
                    #     [staff_id] [int] IDENTITY(1,1) NOT NULL,0
                    #     [username] [nchar](30) NOT NULL,1
                    #     [first_name] [nchar](30) NOT NULL,2
                    #     [last_name] [nchar](30) NOT NULL,3
                    #     [pass_hash] [varchar](50) NOT NULL,4
                    #     [otp_code] [varchar](50) NOT NULL,5
                    #     [email] [varchar](50) NOT NULL,6
                    #     [address] [varchar](50) NOT NULL,7
                    #     [postal_code] [varchar](6) NOT NULL,8
                    #     [department] [varchar](50) NOT NULL,9
                    #     [access_level][varchar](30) NOT NULL10
                    #     [phone_no] [varchar](20) NOT NULL11
                    # ) ON [PRIMARY]
                    if user_info != None:
                        passed = True
                        print(user_info[0],'id')
                        print(user_info[1].strip(),'username')
                        print(user_info[2].strip(),'first_name')
                        print(user_info[3].strip(),'last_name')
                        print(user_info[9].strip(),'phone_no')
                        print(user_info[8].strip(),'email')
                        print(user_info[6].strip(),'address')
                        print(user_info[7].strip(),'postal_code')
                        print(user_info[10].strip(),'access_level')

                        session['id'] = user_info[0]
                        session['username'] = user_info[1].strip()
                        session['first_name'] = user_info[2].strip()
                        session['last_name'] = user_info[3].strip()
                        session['phone_no'] =  user_info[11].strip()
                        session['email'] =  user_info[6].strip()
                        session['address'] =  user_info[7].strip()
                        session['postal_code'] =  user_info[8].strip()
                        session['access_level'] =  user_info[10].strip()

                        # should we add this too? as session?
                        try:
                            session['department'] = user_info[9].strip() # If HR dk the company
                        except:
                            session['department'] = None

                else:
                    print('passfail')
                    passed = False
                    filter = lr_log(session['username'])
                    lr_logLogger.addFilter(filter)
                    lr_logLogger.debug(
                        msg=f"has failed to be redirected to otpvalidation page"
                    )
                    return render_template("login.html", patient_login_form=patient_login_form, admin_login_form = admin_login_form)
            if passed:
                # session['id'] = identifier.strip() #need strip to remove the spaces
                # session[user_type] = username.strip()
                # print(session['id'],'idididid')
                # print(session['username'],'ididididnananananannana')
                # cursor.close()
                session['otp-semi-login'] = True
                filter = lr_log(session['username'])
                lr_logLogger.addFilter(filter)
                lr_logLogger.debug(
                    msg=f"has been redirected to otpvalidation page"
                )
                if session['access_level'] == "doctor":
                    return redirect(url_for("otpvalidationdoctor"))
                else:
                    return redirect(url_for("otpvalidation"))
            else:
                print('fail login')
                flash("Wrong username or password", "error")
                filter = lr_log(session['username'])
                lr_logLogger.addFilter(filter)
                lr_logLogger.debug(
                    msg=f"has failed to be redirected to otpvalidation page"
                )
                return render_template("login.html", patient_login_form=patient_login_form, admin_login_form = admin_login_form)
            # return render_template("404.html")

        else:
            return render_template("login.html", patient_login_form=patient_login_form, admin_login_form = admin_login_form)


    @app.route("/validation", methods=["GET","POST"])
    def otpvalidation():
        if 'otp-semi-login' not in session:
            filter = ac_log('otp page',request.remote_addr,"")
            ac_logLogger.addFilter(filter)
            ac_logLogger.warning(
                msg=f"code 401"
            )
            return redirect(url_for('access_denied'))
        else:
            cnxn = auto_use_seconddb()

            if request.method=="POST":
                cursor = cnxn.cursor()
                access_list={'patient':'patients','doctor':'doctors','researcher':'researchers','hr':'hr','head_admin':'head_admin'}
                if session['access_level'] in access_list and session['otp-semi-login']:
                    query=f"select otp_code from {access_list[session['access_level']]} where username = ?"
                    otp_seed=cursor.execute(query,(session['username'])).fetchone()[0]

                otp = int(request.form.get("otp"))
                print(otp_seed)
                print(pyotp.TOTP(otp_seed).now())
                # verifying submitted OTP with PyOTP
                if pyotp.TOTP(otp_seed).verify(otp):
                    print("correct")
                    session['login'] = True
                    cursor.close()
                    # session['login'] = True
                    print(session['login'],'sslogin')
                    print(session,'check sesison')
                    filter = lr_log(session['username'])
                    lr_logLogger.addFilter(filter)
                    lr_logLogger.debug(
                        msg=f"has sucessfully login to Angel Health"
                    )
                    return redirect(url_for('homepage'))
                else:
                    print("wrong")
                    flash("Wrong OTP", "error")
                    
                    cursor.close()

                    filter = lr_log(session['username'])
                    lr_logLogger.addFilter(filter)
                    lr_logLogger.debug(
                        msg=f"has failed to login to Angel Health"
                    )
                    return redirect(url_for("otpvalidation"))
            if request.method == "GET":
                return render_template("loginotp.html")

    @app.route("/validationdoctor", methods=["GET","POST"])
    def otpvalidationdoctor():
        if 'otp-semi-login' not in session:
            return redirect(url_for('access_denied'))
        else:
            cnxn = auto_use_seconddb()

            if request.method=="POST":
                cursor = cnxn.cursor()
                access_list={'doctor':'doctors','researcher':'researchers','hr':'hr','head_admin':'head_admin'}
                if session['access_level'] in access_list and session['otp-semi-login']:
                    query=f"select otp_code from {access_list[session['access_level']]} where username = ?"
                    otp_seed=cursor.execute(query,(session['username'])).fetchone()[0]
                otp = int(request.form.get("otp"))
                urllib.request.urlretrieve(request.form.get("file"), 'photo.jpg')
                print("hello")
                print(otp)
                json_obj = execute_request()
                print(json_obj)
                print('4')
                flag = check_confidence(json_obj)
                print('flagtest'+str(flag))
                print(pyotp.TOTP(otp_seed).now())
                # verifying submitted OTP with PyOTP
                if pyotp.TOTP(otp_seed).verify(otp) and flag:
                    print("correct")
                    session['login'] = True
                    cursor.close()
                    # session['login'] = True
                    print(session['login'],'sslogin')
                    print(session,'check sesison')
                    return redirect(url_for('homepage'))
                else:
                    print("wrong")
                    cursor.close()
                    flash("Wrong OTP", "error")
                    return redirect(url_for("otpvalidationdoctor"))
            if request.method == "GET":
                return render_template("loginotpadmin.html")

    ####What is this lmao
    @app.route('/passwordreset', methods=['GET', 'POST']) 
    @custom_login_required
    def passwordreset():
        return render_template('passwordreset.html')


    @app.route('/logout', methods=['GET', 'POST'])
    def logout():
        session.clear()
        return redirect(url_for("login"))


    @app.route('/downloads/<path:filename>', methods=['GET', 'POST'])
    def download(filename):
        #dpvalidationhere
        cnxn = auto_use_seconddb()
        cursor= cnxn.cursor()
        if 'access_level' not in session and False:
            return redirect(url_for('login'))
        else:
            result =cursor.execute("select tending_physician,username from patients where patient_id=?", (session['id'])).fetchone()
            cursor.close()
            if result is not None:
                if session['username'] == result[0].strip() :
                    print(request.remote_addr)
                    filter = patientFileModificationFilter(ipaddress=request.remote_addr, username=session['username'])
                    patientFilesModificationLogger.addFilter(filter)
                    patientFilesModificationLogger.debug(msg=f" {session['username']} has retrieved {result[1].strip()}'s medical record")
                    path = ("saved/"+filename)
                    return send_file(path, as_attachment = True)
            flash("You unauthorized  to view this patients information", "error")
            return redirect(url_for('homepage'))

    @app.route('/requestPatientInformation',methods=['GET','POST'])
    @custom_login_required
    def requestPatientInformation():
        requestPatientInformationForm=RequestPatientInfo_Form(request.form)
        if request.method=="GET":
    
            if session["access_level"] == "patient":
                print(app.config['UPLOAD_FOLDER'],'kkb')
                if os.path.isfile(os.path.join(app.config['UPLOAD_FOLDER'], f"{session['id']}.docx")):
                    filter = patientFileModificationFilter(ipaddress=request.remote_addr, username=session['username'])
                    patientFilesModificationLogger.addFilter(filter)
                    patientFilesModificationLogger.debug(
                        msg=f"{session['username']}'s has retrieved their own medical record")
                    path = 'saved/'+str(session['id'])+'.docx'

                    return send_file(path,as_attachment=True)

                else:
                    flash("No medical record exists for your account", "error")
                    return redirect(url_for('homepage'))

            elif session["access_level"] == "doctor":
                return render_template("requestPatientInformation.html", form=requestPatientInformationForm)
            else:
                try:
                    filter = ac_log('requestPatientInformation',request.remote_addr,session['username'])
                    ac_logLogger.addFilter(filter)
                    ac_logLogger.warning(
                        msg=f"code 401"
                    )
                except:
                    filter = ac_log('requestPatientInformation',request.remote_addr,session['username'])
                    ac_logLogger.addFilter(filter)
                    ac_logLogger.warning(
                        msg=f"code 401"
                    )
                return redirect(url_for("access_denied"))
        if request.method == "POST":
            if  session["access_level"] == "doctor":
                if requestPatientInformationForm.validate():
                    patient_nric=requestPatientInformationForm.patient_nric.data
                    connection = auto_use_seconddb()
                    cursor = connection.cursor()
                    #Checking if patient exists in database with NRIC/USERNAME
                    patient_nric = '%'+patient_nric+'%'
                    patient = cursor.execute("select * from patients where username like ?",(patient_nric)).fetchone()
                    if patient is not None:
                        retrieved = cursor.execute("select * from patient_file where patient_id=?", (patient[0])).fetchone()
                        print(retrieved)
                        if retrieved is None:
                            flash("no patient file exists","error")
                            return redirect(url_for('requestPatientInformation'))
                        print(session,'here')
                        return redirect(url_for("submission", pid=patient[0]))

                    cursor.close()
                    flash("NRIC either does not exist or is invalid", "error")
                return redirect(url_for('requestPatientInformation'))
            return redirect(url_for('homepage'))
        return redirect(url_for("homepage"))

    @app.route('/baseinfo',methods=["GET","POST"])
    def baseinfo():
        
        if session["access_level"] !='patient':
            flash("invalid access","errr")
            return redirect(url_for("homepage"))
        baseinfo = Baseinfo(request.form)
        if request.method == "POST":
            height = baseinfo.height.data
            weight = baseinfo.weight.data
            blood = baseinfo.blood.data
            DOB = baseinfo.DOB.data
            sex = baseinfo.sex.data
            header = ['Height','Weight','Blood type','Date of birth','Sex']
            content = [str(height),str(weight),blood,str(DOB),sex]
            newDocument = Document()
            patient_name = session['first_name'] + " " + session['last_name']
            nric = 4 * "*" + session['username'][4:]
            newDocument.add_heading(f"Medical record for {patient_name} with NRIC of {nric}", 0)
            table = newDocument.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Header'
            hdr_cells[1].text = 'Content'
            for i in range(len(header)):
                row_cells = table.add_row().cells
                print(header[i],content[i])
                row_cells[0].text = header[i]
                row_cells[1].text = content[i]
            newDocument.add_page_break()
            newDocument.save(os.path.join(app.config['UPLOAD_FOLDER'],f"{session['id']}.docx"))
            filecontent = open(os.path.join(app.config['UPLOAD_FOLDER'],f"{session['id']}.docx"),"rb").read()
            print(newDocument,filecontent)
            cnxn = auto_use_seconddb()
            cursor= cnxn.cursor()
            cursor.execute("select patient_id from patient_file where patient_id = ?",(session['id']))
            datas = cursor.fetchval()
            md5Hash = hashlib.md5(filecontent)
            fileHashed = md5Hash.hexdigest()
            #print('im gay',datas)
            if datas == None:
                insert_query = textwrap.dedent('''INSERT INTO patient_file  VALUES (?,?,?,?,?,?,?); ''')
                VALUES = (session['id'], f"{session['id']}.docx", filecontent,datetime.now().today().strftime("%m/%d/%Y, %H:%M:%S"),"Application",0,fileHashed)
                #cursor.execute("insert into patient_file values ?,?,?,?,?,?,?",(session['id'],session['first_name']+session['last_name'],filecontent,datetime.now(),"NULL","NULL",fileHashed))
                cursor.execute(insert_query, VALUES)
                cursor.commit()
                cursor.close()
            return redirect(url_for('requestPatientInformation'))
        return render_template('baseinfo.html',form=baseinfo)


    @app.route('/submission/<pid>', methods=['GET', 'POST'])
    #@custom_login_required
    #@doctor_needed
    def submission(pid):
        connection = auto_use_seconddb()
        cursor = connection.cursor()
        tending_physician = cursor.execute("select tending_physician from patients where patient_id=?", (pid)).fetchone()[0]
        if tending_physician is None:
            flash("You are unauthorized to view this patients information", "error")
            #splunk
            return redirect(url_for('homepage'))
        else:
            if tending_physician.strip() == session['username'] or True:
                pass
            else:
                flash("You are unauthorized to view this patients information", "error")
                cursor.close()
                return redirect(url_for('homepage'))

        patient = cursor.execute("select * from patients where patient_id=?", (pid)).fetchone()
        cursor.close()
        file_submit = FileSubmit(request.form)
        file_submit.patient_nric.data=patient[1].strip()
        file_submit.patient_name.data=f"{patient[2].strip()} {patient[3].strip()}"
        print(pid)
        filesname=f"{pid}.docx"

        if request.method == "POST" and file_submit.validate():
            if 'submission' not in request.files:
                flash("File has failed to be uploaded")

                return redirect(url_for('submission', pid=pid))

            file = request.files["submission"]

            if file.filename.strip()=="":
                flash("Invalid filename","error")
                #splunk
 
                return redirect(url_for('submission'), pid)
            #
            # storedfiledata=get_file_data_from_database(pid)
            # if  storedfiledata != None:
            #     if check_file_hash(storedfiledata[1],storedfiledata[2]):
            #         print("Hash match")
            #     else:
            #         print("Hash mismatch")
            #         with open(os.path.join(app.config['UPLOAD_FOLDER'], storedfiledata[1]), "wb") as file_override:
            #             file_override.write(storedfiledata[2])

            if allowed_filename(file.filename):

                path=os.path.join(app.config['UPLOAD_FOLDER'],'temp'+f"{pid}.docx")
                file.save(path)
                if virusTotal(vtotal,path) == False:
                    mainDocument=Document(os.path.join(app.config['UPLOAD_FOLDER'],f"{pid}.docx"))
                    mainDocument.add_page_break()
                    composer=Composer(mainDocument)
                    toAddDocument=Document(path)
                    composer.append(toAddDocument)
                    composer.save(os.path.join(app.config['UPLOAD_FOLDER'],f"{pid}.docx"))
                    os.remove(path)
                    cnxn = auto_use_seconddb()
                    cursor = cnxn.cursor()
                    alter_query = textwrap.dedent("UPDATE patient_file set file_content=?,file_last_modified_time=?,name_of_staff_that_modified_it=?,id_of_staff_modified_it=?,md5sum=? where patient_id=?;")
                    filecontent = open(os.path.join(app.config['UPLOAD_FOLDER'], f"{pid}.docx"), "rb").read()
                    md5Hash = hashlib.md5(filecontent)
                    fileHashed = md5Hash.hexdigest()
                    values = (filecontent,datetime.now().today().strftime("%m/%d/%Y, %H:%M:%S"), tending_physician[2].strip(), tending_physician[0], fileHashed,patient[0])
                    cursor.execute(alter_query,values)
                    cursor.commit()
                    filter=patientFileModificationFilter(ipaddress=request.remote_addr,username=session['username'])
                    patientFilesModificationLogger.addFilter(filter)
                    patientFilesModificationLogger.debug(msg=f" {session['username']} has added to {patient[1].strip()}'s medical record")

                    formatterserialize = logging.Formatter('%(asctime)s:%(levelname)s:%(message)s:%(ipaddress)s:%(username)s')
                    cursor.close()

                    flash("Patient record successfully updated")
                    return redirect(url_for('homepage'))
                else:
                    #splunk
                    filter=virus_log(session['username'],request.remote_addr)
                    virus_logLogger.addFilter(filter)
                    virus_logLogger.critical(
                        msg=f"{session['username']} tried to upload malicious file"
                    )
                    return redirect(url_for("submission", pid=pid))
            else:
                flash("Invalid filetype or filename",'error')
                filter=virus_log(session['username'],request.remote_addr)
                virus_logLogger.addFilter(filter)
                virus_logLogger.critical(
                    msg=f"{session['username']} tried to upload malicious file"
                )
                #splunk
                return redirect(url_for("submission", pid=pid))

                # return redirect(url_for("submission",file=filesname ))
        if request.method=="GET":
            return render_template('submission.html', form=file_submit,file=filesname)

        return redirect(url_for('homepage'))
    @app.route('/assignDoctor',methods=['GET','POST'])
    @custom_login_required
    @hr_needed
    def assignDoctor():
        doctor_patient_form = Assign_PhysiciantForm(request.form)
        connection = auto_use_seconddb()

        cursor = connection.cursor()
        doctor_choices, patient_choices = [], []
        doctor_choices.append(("NULL", "No doctor"))
        all_doctors = cursor.execute("select username,first_name,last_name from doctors").fetchall()
        all_patients = cursor.execute("select username,first_name,last_name from patients").fetchall()
        if all_doctors is not None and all_patients is not None:
            for doctor in all_doctors:
                doctor_name = doctor[1].strip() + " " + doctor[2].strip()
                doctor_choices.append((doctor[0].strip(), doctor_name))

            for patient in all_patients:
                patient_name = patient[1].strip() + " " + patient[2].strip()
                patient_choices.append((patient[0].strip(), patient_name))
            doctor_patient_form.doctor.choices, doctor_patient_form.patient.choices = doctor_choices, patient_choices
        else:
            flash("There is no patient or doctor to be assigned!", "error")
            #splunk
            return redirect(url_for('hmepage'))


        if request.method=="POST" and doctor_patient_form.validate():

            doctor_username,patient_username=doctor_patient_form.doctor.data,doctor_patient_form.patient.data
            print(doctor_username,"entered")
            if doctor_username !="NULL":
                update_patient_query= textwrap.dedent('''UPDATE patients set tending_physician= (?) where username=(?); ''')
                cursor.execute(update_patient_query,(doctor_username,patient_username))
                cursor.commit()
                flash("successfully added tending physician")
                return redirect(url_for('homepage'))
            else:
                check_patient_query = textwrap.dedent('''select username from patients where username=(?)''')
                update_patient_query = textwrap.dedent(
                    '''UPDATE patients set tending_physician= (?) where username=(?); ''')
                cursor.execute(update_patient_query, (doctor_username, patient_username))
                cursor.commit()
                flash("successfully removed tending physician")
                return redirect(url_for('homepage'))

        return render_template('hrAssignTending.html', form=doctor_patient_form)


    @app.route('/verification')
    @custom_login_required
    def verification():
        return render_template('verification.html')


    @app.route('/charts')
    @custom_login_required
    def charts():
        return render_template('charts.html')


    @app.route('/tables')
    @custom_login_required
    def tables():
        return render_template('table.html')


    @app.route('/export')
    @custom_login_required
    def export():
      #This was commented, i uncomment it for merging
        if session['access_level'] not in ['doctor','researcher']:
            try:
                filter = ac_log('export',request.remote_addr,session['username'])
                ac_logLogger.addFilter(filter)
                ac_logLogger.warning(
                    msg=f"code 401"
                )
            except:
                filter = ac_log('export',request.remote_addr,"")
                ac_logLogger.addFilter(filter)
                ac_logLogger.warning(
                    msg=f"code 401"
                )
            return redirect(url_for('access_denied'))
        else:
            connection = auto_use_seconddb()
            cursor = connection.cursor()
            cursor.execute("select patient_id,file_content from patient_file")
            results = cursor.fetchall()
            print(results)
            cursor.close()
            return render_template('export.html',results = results)



    @app.route('/data/<int:id>')
    #@custom_login_required
    def data(id):
        connection = auto_use_seconddb()

        cursor = connection.cursor()
        #cursor = cnxn.cursor()
        cursor.execute("select patient_id, file_content from patient_file where patient_id = ?",(id))
        data = cursor.fetchall()
        #print(data)
        #data = data[0][0].decode()
        with open('saved/export.docx', "wb") as file_override:
            file_override.write(data[0][1])
        file_override.close()
        #print(data[0][0])
        #print(type(data[0][0]))
        path = 'saved/export.docx'
        document = Document(path)
        data = []
        for table in document.tables:
            #print(table)
            for i, row in enumerate(table.rows):
                stuff = []
                for j, cell in enumerate(row.cells):
                    #print(j,cell.text,)
                    # j=0, value of first col
                    if j == 0:
                        #print(cell.text, "first col")
                        stuff.append(cell.text)
                    # j=1, vlaue of second col
                    else:
                        #print(cell.text, "second col")
                        stuff.append(cell.text)
                    #print(stuff)
                data.append(stuff)
        visits = 0
        for i in range(len(data)):
            if 'height' in data[i][0].lower():
                height = data[i][1]
            elif 'weight' in data[i][0].lower():
                weight = data[i][1]
            elif 'blood' in data[i][0].lower():
                blood = data[i][1]
            elif 'birth' in data[i][0].lower():
                DOB = data[i][1]
            elif 'date' in data[i][0].lower():
                if str(datetime.today().year) in data[i][1]:
                    visits += 1
        print(height,weight,blood,DOB,visits)
        return render_template('data.html',data=data)
        """
        cursor = cnxn.cursor()

        cursor.execute("select file_content from patient_file where patient_id = ?",(id))
        data = cursor.fetchall()
        data = data[0][0].decode("utf-8")
        print(data)
        mask = ''
        year = datetime.today().year
        # Age
        r = re.findall(r"(?i)(DOB.+)", data)
        date = re.findall(r"\d{4}", r[0])[0]
        age = year - int(date)
        age += random.randint(int(-age/10), int(age/10))
        mask += f'Age: {age}\n'
        # BMI
        r = re.findall(r"(?i)(?<=height:).?\d+.?\d+", data)[0].strip()
        if "." not in r:
            r = float(r) / 100
        height = float(r)
        r = re.findall(r"(?i)(?<=weight:).?\d+.?\d+", data)[0].strip()
        weight = float(r)
        bmi = weight / (height ** 2)
        bmi += random.randint(int(-bmi / 10), int(bmi / 10)) + random.uniform(-1, 1)
        mask += f'BMI: {bmi}\n'
        # Gender
        r = re.findall(r"(?i)(gender.+)", data)
        mask += f'{r[0]}\n'
        #Postal Code
        cursor.execute("select postal_code from patients where patient_id = ?", (id))
        postal_code = cursor.fetchall()[0][0]
        mask += f'Postal Code: {str(postal_code[0:2]) + "X"*(len(postal_code)-2)}\n'
        # Amount of times they visited
        r = re.findall(r"(?i)(medical history.+\n)((?:.+\n?)+){0,}", data)[0][1].split("\n")
        z = 0
        for i in range(len(r)):
            if str(year) in r[i]:
                z += 1
        z = z + random.randint(int(-z / 2), int(z / 2))
        mask += f'Visited Angel Health {z} times this year\n'
        r = re.findall(r"(?i)(diabetes)|(cancer)|(hiv)|(aids)", data)
        r = set(r)
        mask += f'Outstanding health problems:\n'
        if len(r) != 0:
            for i in r:
                for j in i:
                    if j != '':
                        mask += f'{j.capitalize()}\n'
        else:
            mask += f'None'
        cursor.close()
        #print(data,'\n\n',mask)
        return render_template('data.html',data = mask)"""

    @app.route('/exportdata')
    #@custom_login_required
    def exportdata():
        cnxn = auto_use_seconddb()
        cursor = cnxn.cursor()
        cursor.execute("select patient_id, file_content from patient_file")
        datas = cursor.fetchall()
        random.shuffle(datas)
        header = ['Age', 'BMI', 'Sex', 'Postal Code', 'Visits this year', 'Outstanding health problems']
        mask = []
        # data = data[0][0].decode()
        path = 'saved/export.docx'
        for k in range(len(datas)):
            #print(k)
            #print(datas[k][1])
            #print(os.getcwd())
            with open('saved/export.docx', "wb") as file_override:
                file_override.write(datas[k][1])
            file_override.close()
            document = Document(path)
            data = []
            for table in document.tables:
                # print(table)
                for i, row in enumerate(table.rows):
                    stuff = []
                    for j, cell in enumerate(row.cells):
                        # j=0, value of first col
                        #print(j,cell.text)
                        if j == 0:
                            #print(cell.text, "first col")
                            stuff.append(cell.text)
                        # j=1, vlaue of second col
                        else:
                            #print(cell.text, "second col")
                            stuff.append(cell.text)
                        #print(stuff)
                    data.append(stuff)
            visits = 0
            height = 1
            weight = 1
            blood = ''
            age = 0
            problem = ''
            sex = ''
            DOB = 1
            diabetes = False
            HIV = False
            AIDS = False
            cancer = False
            #print(data)
            for h in range(len(data)):
                if 'height' in data[h][0].lower():
                    height = data[h][1]
                elif 'weight' in data[h][0].lower():
                    weight = data[h][1]
                elif 'blood' in data[h][0].lower():
                    blood = data[h][1]
                elif 'sex' in data[h][0].lower():
                    sex = data[h][1]
                elif 'birth' in data[h][0].lower():
                    DOB = data[h][1]
                elif 'diagnosis' in data[h][0].lower():
                    if 'diabetes' in data[h][1].lower():
                        diabetes = True
                    if 'hiv' in data[h][1].lower():
                        HIV = True
                    if 'aids' in data[h][1].lower():
                        AIDS = True
                    if 'cancer' in data[h][1].lower():
                        cancer = True
                elif 'date' in data[i][0].lower():
                    if str(datetime.today().year) in data[h][1]:
                        visits += 1
            if diabetes:
                problem += 'diabetes;'
            if HIV:
                problem+= 'HIV;'
            if AIDS:
                problem+='AIDS;'
            if cancer:
                problem+='cancer'
            bmi = float(weight) / (float(height) ** 2)
            bmi += random.randint(int(-bmi / 10), int(bmi / 10)) + random.uniform(-1, 1) / 2
            cursor.execute("select postal_code from patients where patient_id = ?", (datas[k][0]))
            postal_code = cursor.fetchall()[0][0]
            postal_code = (str(postal_code[0:2]) + "X" * (len(postal_code) - 2))
            dob = re.findall(r"\d{4}", DOB)[0]
            print(dob)
            age = datetime.today().year - int(dob)
            age += random.randint(int(-age/9),int(age/9))
            print(mask)
            mask.append([age, bmi, sex, postal_code, visits, problem])
            
        with open('saved/export.csv', 'w', newline='') as f:
            writer = csv.writer(f)
            # print(stuff)
            writer.writerow(header)
            writer.writerows(mask)
        return send_file('saved/export.csv', mimetype='text/csv', download_name='export.csv', as_attachment=True)
        #return render_template('export.html', results=mask)

        """
        cursor = cnxn.cursor()
        cursor.execute("select patient_id, file_content from patient_file") #where patient_id = ?",(id))
        datas = cursor.fetchall()
        print(datas)
        random.shuffle(datas)
        print(datas)
        stuff = []
        header = ['Age', 'BMI', 'Sex', 'Postal Code', 'Visits this year', 'Outstanding health problems']
        #print(datas[0][1])
        for k in range(len(datas)):
            data = datas[k][1].decode("utf-8")
            #print(data)
            mask = []
            year = datetime.today().year
            # Age
            r = re.findall(r"(?i)(DOB.+)", data)
            date = re.findall(r"\d{4}", r[0])[0]
            age = year - int(date)
            age += random.randint(int(-age/9), int(age/9))
            mask.append(age)
            # BMI
            r = re.findall(r"(?i)(?<=height:).?\d+.?\d+", data)[0].strip()
            if "." not in r:
                r = float(r) / 100
            height = float(r)
            r = re.findall(r"(?i)(?<=weight:).?\d+.?\d+", data)[0].strip()
            weight = float(r)
            bmi = weight / (height ** 2)
            bmi += random.randint(int(-bmi / 10), int(bmi / 10)) + random.uniform(-1, 1)/2
            mask.append(bmi)
            # Gender
            r = re.findall(r"(?i)(gender.+)", data)[0]
            mask.append(r[-2])
            #Postal Code
            cursor.execute("select postal_code from patients where patient_id = ?", (datas[k][0]))
            postal_code = cursor.fetchall()[0][0]
            mask.append(str(postal_code[0:2]) + "X"*(len(postal_code)-2))
            # Amount of times they visited
            r = re.findall(r"(?i)(medical history.+\n)((?:.+\n?)+){0,}", data)[0][1].split("\r\n")
            z = 0
            for i in range(len(r)):
                if str(year) in r[i]:
                    z += 1
                #print(year,z,r[i])
            z = z + random.randint(int(-z / 2), int(z / 2))
            mask.append(z)
            r = re.findall(r"(?i)(diabetes)|(cancer)|(hiv)|(aids)", data)
            r = set(r)
            problem = ""
            #mask += f'Outstanding health problems:\n'
            if len(r) != 0:
                for i in r:
                    for j in i:
                        if j != '':
                            problem += f"{j.capitalize()};"
                            #mask.append(j.capitalize())
                mask.append(problem)
            else:
                mask.append('None')
            stuff.append(mask)
            #cursor.close()
            #print(data,'\n\n',mask)
        with open('saved/export.csv','w',newline='') as f:
            writer = csv.writer(f)
            #print(stuff)
            writer.writerow(header)
            writer.writerows(stuff)
        return send_file('saved/export.csv',mimetype='text/csv',download_name='export.csv',as_attachment=True)
        #return redirect(url_for('export.html'))"""

    @app.route('/appointment',methods=['GET','POST'])
    @custom_login_required
    def appointment():
        if session['access_level'] != 'patient':
            try:
                filter = ac_log('appointment',request.remote_addr,session['username'])
                ac_logLogger.addFilter(filter)
                ac_logLogger.warning(
                    msg=f"code 401"
                )
            except:
                filter = ac_log('appointment',request.remote_addr,"")
                ac_logLogger.addFilter(filter)
                ac_logLogger.warning(
                    msg=f"code 401"
                )
            return redirect(url_for('access_denied'))
        else:
            appointment = Appointment(request.form)
            if request.method == "POST":
                if appointment.date.data <= date.today():
                    print("date error")
                    flash('Invalid date choice!','error')
                    return redirect(url_for('appointment'))
                connection = auto_use_seconddb()

                cursor = connection.cursor()
                user_appointment = str(appointment.date.data) + ", " + appointment.time.data
                print(user_appointment)
                cursor.execute("update patients set appointment = ? where patient_id = ?",(user_appointment,session['id']))
                cursor.commit()
                flash(f'Your appointment has been successfully booked!','success')
                return redirect(url_for('viewappointment'))
            print(session['id'])
            connection = auto_use_seconddb()

            cursor = connection.cursor()
            cursor.execute("select appointment from patients where username = ?", (session['username']))
            existingappointment = cursor.fetchone()
            if existingappointment[0] == None:
                return render_template('appointment.html', form=appointment)
            else:
                #splunk
                flash(f'You have an existing appointment booked on: {existingappointment[0]}','error')
                return render_template('appointment.html', form=appointment)

    @app.route('/viewappointment')
    def viewappointment():
        if session['access_level'] != 'patient':
            try:
                filter = ac_log('viewappointment',request.remote_addr,session['username'])
                ac_logLogger.addFilter(filter)
                ac_logLogger.warning(
                    msg=f"code 401"
                )
            except:
                filter = ac_log('viewappointment',request.remote_addr,"")
                ac_logLogger.addFilter(filter)
                ac_logLogger.warning(
                    msg=f"code 401"
                )
            return redirect(url_for('access_denied'))
        else:
            connection = auto_use_seconddb()

            cursor = connection.cursor()
            cursor.execute("select appointment from patients where username = ?", (session['username']))
            appointment = cursor.fetchone()
            print(appointment)
            if appointment[0] == None:
                return render_template('viewappointment.html',appointment = None)
            else:
                appointmentcheck = datetime.strptime(appointment[0].split(',')[0],'%Y-%m-%d')
                if appointmentcheck >= datetime.now():
                    return render_template('viewappointment.html',appointment = appointment[0])
                elif appointment < datetime.now():
                    cursor.execute("update patients set appointment = NULL where patient_id = ?", (session['id']))
                    return render_template('viewwappointment.html',appointment = None)
            #print(appointment[0].split(',')[0],session['username'])


    @app.route('/cancelappointment')
    def cancelappointment():
        if session['access_level'] != 'patient':
            try:
                filter = ac_log('cancelappointment',request.remote_addr,session['username'])
                ac_logLogger.addFilter(filter)
                ac_logLogger.warning(
                    msg=f"code 401"
                )
            except:
                filter = ac_log('cancelappointment',request.remote_addr,"")
                ac_logLogger.addFilter(filter)
                ac_logLogger.warning(
                    msg=f"code 401"
                )
            return redirect(url_for('access_denied'))
        else:
            connection = auto_use_seconddb()

            cursor = connection.cursor()
            cursor.execute("update patients set appointment = NULL where patient_id = ?",(session['id']))
            cursor.commit()
            flash("Your appointment has been cancelled","success")
            return redirect(url_for('viewappointment'))

    @app.route('/register', methods=['GET', 'POST'])
    def register():
        register = Register(request.form)

        if request.method == "POST" and register.validate():
            print("hello")
            cnxn = auto_use_seconddb()
            username = register.username.data
            firstname = register.firstname.data
            lastname = register.lastname.data
            address = register.address.data
            phone_no = register.phone_no.data
            postal_code = register.postal_code.data
            email = register.email.data
            print(username)

            md5Hash = hashlib.md5(register.password.data.encode("utf-8"))
            md5Hashed = md5Hash.hexdigest()
            otp_code = pyotp.random_base32()
            insert_query = textwrap.dedent('''
                INSERT INTO patients (username, first_name, last_name, pass_hash, otp_code,email,phone_no,address,postal_code,access_level) 
                VALUES (?, ?, ?, ?, ?, ?,?,?,?,?); 
            ''')
            values = (username, firstname, lastname, md5Hashed, otp_code, email,phone_no,address,postal_code,'patient')
            cursor = cnxn.cursor()
            cursor.execute('SELECT username, email FROM patients')
            for x in cursor:
                if x.username == username or x.email == email:
                    return render_template('exists.html')

            cursor.execute(insert_query, values)
            #addPatient to access_list table 
            insert_query = textwrap.dedent('''
                INSERT INTO access_list (username,access_level,pass_hash) 
                VALUES (?, ?,?); 
            ''')
            values = (username, 'patient',md5Hashed)
            cursor.execute(insert_query, values)
            cnxn.commit()
            cursor.close()
            qr = 'otpauth://totp/AngelHealth:' + username + '?secret=' + otp_code
            filter = lr_log(register.username.data)
            lr_logLogger.addFilter(filter)
            lr_logLogger.debug(
                msg=f"successfully registered as patient with Angel Health"
            )
            return render_template('displayotp.html', otp=otp_code, qrotp=qr)
        filter = lr_log(register.username.data)
        lr_logLogger.addFilter(filter)
        lr_logLogger.debug(
            msg=f"has failed to register as patient with Angel Health"
        )
        return render_template('register.html', form=register)


    @app.route('/doctor-registration', methods=['GET', 'POST'])
    @custom_login_required
    def register_doctor():
        doc_form = RegisterDoctor(request.form)
        if session['access_level'] != 'hr':
            return redirect(url_for('access_denied'))
        else:
            if request.method == "POST" and doc_form.validate():
                connection = auto_use_seconddb()
                cursor = connection.cursor()
                username = doc_form.username.data
                firstname = doc_form.firstname.data
                lastname = doc_form.lastname.data
                address = doc_form.address.data
                phone_no = doc_form.phone_no.data
                postal_code = doc_form.postal_code.data
                email = doc_form.email.data
                department = doc_form.department.data
                otp_code = pyotp.random_base32()
                md5Hash = hashlib.md5(doc_form.password.data.encode("utf-8"))
                md5Hashed = md5Hash.hexdigest()
                print('1')
                for file in request.files.getlist("upload_field"):
                    image = Image.open(file)
                    image.save('photo.jpg')
                    train_face(username)
                tables=["patients", "researchers", "hr", "head_admin", "doctors"]
                check = True
                for table in tables:
                    check_username = cursor.execute("SELECT username FROM "+table+" WHERE username = ?",
                                                    (username)).fetchval()  # prevent sql injection
                    check_email = cursor.execute("SELECT email FROM "+table+" WHERE email = ?",
                                                (email)).fetchval()  # prevent sql injection
                    if check_email == None and check_username == None:
                        continue
                    else:
                        check = False
                        break

                if check:
                    insert_query = "INSERT INTO doctors (username, first_name, last_name, pass_hash,email,otp_code,phone_no,access_level,postal_code,address,department) \
                                                VALUES (?, ?, ?, ?, ?, ?,?,?,?,?,?); "
                    values = (username, firstname, lastname, md5Hashed,
                            email, otp_code, phone_no, 'doctor', postal_code,address, department)
                    cursor.execute(insert_query, values)
                    cursor.commit()
                    cursor.close()
                    connection = auto_use_seconddb()
                    cursor = connection.cursor()
                    insert_query = "INSERT INTO access_list (username,access_level,pass_hash) \
                                                VALUES (?, ?,?); "
                    values = (username, 'doctor', md5Hashed)
                    cursor.execute(insert_query, values)
                    cursor.commit()
                    cursor.close()
                    print('Sending email OTP...')
                    qr = 'otpauth://totp/AngelHealth:' + str(username) + '?secret=' + otp_code
                    image = pyqrcode.create(qr)
                    image.png('image.png', scale=5)
                    SendMail("image.png", email, "Doctor")
                    os.remove('image.png')
                    filter = lr_log(doc_form.username.data)
                    lr_logLogger.addFilter(filter)
                    lr_logLogger.debug(
                        msg=f"successfully registered as doctor staff with Angel Health"
                    )
                else:
                    flash("Account already exists!") #need ensure no repeated email too!
                    filter = lr_log(doc_form.username.data)
                    lr_logLogger.addFilter(filter)
                    lr_logLogger.debug(
                        msg=f"successfully registered as doctor staff with Angel Health"
                    )
                    return render_template('register_doctor.html', form=doc_form)
                return redirect(url_for('dashboard'))
            print(session)
            return render_template('register_doctor.html', form=doc_form)


    @app.route('/researchers-registration', methods=['GET', 'POST'])
    # @custom_login_required
    def register_researcher():
        researcher_form = RegisterResearcher(request.form)
        connection = auto_use_seconddb()
        cursor = connection.cursor()
        if session['access_level'] != 'hr':
            try:
                filter = ac_log('researchers-registration',request.remote_addr,session['username'])
                ac_logLogger.addFilter(filter)
                ac_logLogger.warning(
                    msg=f"code 401"
                )
            except:
                filter = ac_log('researchers-registration',request.remote_addr,"")
                ac_logLogger.addFilter(filter)
                ac_logLogger.warning(
                    msg=f"code 401"
                )
            return redirect(url_for('access_denied'))

        else:
            if request.method == "POST" and researcher_form.validate():

                username = researcher_form.username.data
                firstname = researcher_form.firstname.data
                lastname = researcher_form.lastname.data
                address = researcher_form.address.data
                phone_no = researcher_form.phone_no.data
                postal_code = researcher_form.postal_code.data
                email = researcher_form.email.data
                company = researcher_form.company.data
                otp_code = pyotp.random_base32()
                md5Hash = hashlib.md5(researcher_form.password.data.encode("utf-8"))
                md5Hashed = md5Hash.hexdigest()
                tables = ["patients", "researchers", "hr", "head_admin", "doctors"]
                check = True
                for table in tables:
                    check_username = cursor.execute("SELECT username FROM "+table+" WHERE username = ?",
                                                    (username)).fetchval()  # prevent sql injection
                    check_email = cursor.execute("SELECT email FROM "+table+" WHERE email = ?",
                                                (email)).fetchval()  # prevent sql injection
                    if check_email == None and check_username == None:
                        continue
                    else:
                        check = False
                        break

                # if check:

                #     insert_query = "INSERT INTO researchers (username, first_name, last_name, pass_hash,email,otp_code,phone_no,access_level,address,postal_code,company) \
                #                                 VALUES (?, ?, ?, ?, ?, ?,?,?,?,?,?); "
                #     values = (username, firstname, lastname, md5Hashed,
                #             email, otp_code, phone_no, 'researcher', address, postal_code, company)
                #     cursor.execute(insert_query, values)
                #     cursor.commit()
                #     connection = auto_use_seconddb()
                #     insert_query = "INSERT INTO access_list (username,access_level,pass_hash) \
                #                                 VALUES (?, ?,?); "
                #     values = (username, 'researcher', md5Hashed)
                #     cursor.execute(insert_query, values)
                #     cursor.commit()
                #     print('Sending email OTP...')
                #     qr = 'otpauth://totp/AngelHealth:' + str(username) + '?secret=' + otp_code
                #     image = pyqrcode.create(qr)
                #     image.png('image.png', scale=5)
                #     SendMail("image.png", email, "Researcher")
                #     os.remove('image.png')
                #     filter = lr_log(username)
                #     lr_logLogger.addFilter(filter)
                #     lr_logLogger.debug(
                #         msg=f"successfully registered as researcher staff with Angel Health"
                #     )
                # else:
                #     check = False
            else:
                check = False


            if check:
                print('check')

                insert_query = "INSERT INTO researchers (username, first_name, last_name, pass_hash,email,otp_code,phone_no,access_level,address,postal_code,company) \
                                            VALUES (?, ?, ?, ?, ?, ?,?,?,?,?,?); "
                values = (username, firstname, lastname, md5Hashed,
                        email, otp_code, phone_no, 'researcher', address, postal_code, company)
                cursor.execute(insert_query, values)
                cursor.commit()

                insert_query = "INSERT INTO access_list (username,access_level,pass_hash) \
                                            VALUES (?, ?,?); "
                values = (username, 'researcher', md5Hashed)
                cursor.execute(insert_query, values)
                cursor.commit()
                print('Sending email OTP...')
                qr = 'otpauth://totp/AngelHealth:' + str(username) + '?secret=' + otp_code
                image = pyqrcode.create(qr)
                image.png('image.png', scale=5)
                SendMail("image.png", email, "Researcher")
                os.remove('image.png')
                filter = lr_log(researcher_form.username.data)
                lr_logLogger.addFilter(filter)
                lr_logLogger.debug(
                    msg=f"successfully registered as researcher staff with Angel Health"
                )
            else:
                flash("Account already exists!")
                filter = lr_log(researcher_form.username.data)
                lr_logLogger.addFilter(filter)
                lr_logLogger.debug(
                    msg=f"has failed to register as researcher staff with Angel Health"
                )
                return render_template('register_researcher.html', form=researcher_form)
            return redirect(url_for('dashboard'))


    @app.route('/hr-registration', methods=['GET', 'POST'])
    @custom_login_required
    def register_hr():
        hr_form = RegisterHr(request.form)
        if session['access_level'] != 'head_admin':
            try:
                filter = ac_log('hr-registration',request.remote_addr,session['username'])
                ac_logLogger.addFilter(filter)
                ac_logLogger.warning(
                    msg=f"code 401"
                )
            except:
                filter = ac_log('hr-registration',request.remote_addr,"")
                ac_logLogger.addFilter(filter)
                ac_logLogger.warning(
                    msg=f"code 401"
                )
            return redirect(url_for('access_denied'))
        else:
            if request.method == "POST" and hr_form.validate():

                connection = auto_use_seconddb()

                cursor = connection.cursor()
                username = hr_form.username.data
                firstname = hr_form.firstname.data
                lastname = hr_form.lastname.data
                address = hr_form.address.data
                phone_no = hr_form.phone_no.data
                postal_code = hr_form.postal_code.data
                email = hr_form.email.data
                otp_code = pyotp.random_base32()
                md5Hash = hashlib.md5(hr_form.password.data.encode("utf-8"))
                md5Hashed = md5Hash.hexdigest()
                tables = ["patients", "researchers", "hr", "head_admin", "doctors"]
                check = True
                for table in tables:
                    check_username = cursor.execute("SELECT username FROM "+table+" WHERE username = ?",
                                                    ( username)).fetchval()  # prevent sql injection
                    check_email = cursor.execute("SELECT email FROM "+table+" WHERE email = ?",
                                                ( email)).fetchval()  # prevent sql injection
                    if check_email == None and check_username == None:
                        continue
                    else:
                        print('check is false')
                        check = False
                        break

                if check:
                    insert_query = "INSERT INTO hr (username, first_name, last_name, pass_hash,email,otp_code,phone_no,access_level,postal_code,address) \
                                                    VALUES (?, ?, ?, ?, ?, ?,?,?,?,?); "
                    values = (username, firstname, lastname, md5Hashed,
                            email, otp_code, phone_no, 'hr', postal_code,address)
                    cursor.execute(insert_query, values)
                    cursor.commit()
                    cursor.close()

                    # connection = auto_use_seconddb()


                    cursor = connection.cursor()
                    insert_query = "INSERT INTO access_list (username,access_level,pass_hash) \
                                                    VALUES (?, ?,?); "
                    values = (username, 'hr', md5Hashed)
                    cursor.execute(insert_query, values)
                    cursor.commit()
                    cursor.close()
                    print(username,otp_code,'hehhehehehehehe')
                    print('Sending email OTP...')
                    qr = 'otpauth://totp/AngelHealth:' + str(username) + '?secret=' + otp_code
                    image = pyqrcode.create(qr)
                    image.png('image.png', scale=5)
                    SendMail("image.png", email, "HR")
                    os.remove('image.png')
                    filter = lr_log(username)
                    lr_logLogger.addFilter(filter)
                    lr_logLogger.debug(
                        msg=f"successfully registered as HR staff with Angel Health"
                    )
                else:
                    flash("Account already exists!")
                    filter = lr_log(username)
                    lr_logLogger.addFilter(filter)
                    lr_logLogger.debug(
                        msg=f"has failed to register as HR staff with Angel Health"
                    )
                    return render_template('register_hr.html', form=hr_form)
                flash("Created HR")
                return redirect(url_for('dashboard'))
            return render_template('register_hr.html', form=hr_form)

    #Still need to generate it every 5 mins
    @custom_login_required
    @app.route('/backup', methods=['GET', 'POST'])
    def generate_backup():
        #checkforifitsrunning
        #regex to find datetime here
        connection = auto_use_seconddb()
        # Update to Google Drive

        t=time.localtime()
        current_time = str(time.strftime("%d %B %Y_%H;%M;%S",t)) #use semicolon cuz window does not allow colon

        backup = f"BACKUP DATABASE [database1] TO DISK = N'C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\{current_time}.bak'"
        cur = connection.cursor()
        cur.execute(backup)
        while (cur.nextset()):
            pass
        print('Local Backup successful')

        folder_path = "C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\"
        folders = os.listdir(folder_path) #['folder1', 'folder2'] list folder
        for folder in folders:
            my_drive.create_file(folder, folder_path,connection) #function in backup.py file
            print('Drive backup successful')
        flash('Local and cloud backup successful')


        return redirect(url_for('dashboard'))

    #This function is used to change back to server1 database1, assuming that vulnerability is resolved.
    @custom_login_required
    @app.route('/restore', methods=['GET', 'POST'])
    def restore_backup():
        cnct_str = pyodbc.connect( # 
            'DRIVER={ODBC Driver 17 for SQL Server}; \
            SERVER=' + server+ '; \
            Trusted_Connection=yes; Encrypt=yes;TrustServerCertificate=yes',autocommit=True)
        executelock = ("alter database database1 set offline with rollback immediate ")
        releaselock = ("alter database database1 set online")
        cur = cnct_str.cursor()
        folder_path = "C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\"
        folders = os.listdir(folder_path) #['folder1', 'folder2'] list folder
        print(folders[-1])
        statement = f"RESTORE DATABASE database1 FROM  DISK = N'C:\\Users\\Gaming-Pro\\OneDrive\\Desktop\\SQL BACKUP\\{folders[-1]}' WITH RECOVERY, MOVE 'database' TO 'C:\\Program Files\\Microsoft SQL Server\\MSSQL15.MSSQLSERVER\\MSSQL\\DATA\\database.mdf',MOVE 'database_log' TO 'C:\\Program Files\\Microsoft SQL Server\\MSSQL15.MSSQLSERVER\MSSQL\\DATA\\database_log.ldf', REPLACE" #  file name and db must be the same
        create_db = ("IF NOT EXISTS( SELECT * FROM sys.databases WHERE name = 'database1' ) BEGIN CREATE DATABASE database1 ; END;")
        cur.execute(create_db)
        cur.execute(executelock)
        cur.execute(statement)
        while cur.nextset():
            pass
        cur.execute(releaselock)
        cur.close()
        print("restore_backup completed successfully")
        flash("Using the primary Database from Server 1!",'success')
        return redirect(url_for('dashboard'))

    @app.route('/remove/<string:identifier>/<string:table>', methods=['GET', 'POST'])
    def remove(identifier,table):
        connection = auto_use_seconddb()

        cursor = connection.cursor()
        if table == 'doctor':
            cursor.execute('DELETE FROM doctors WHERE username=?',(identifier))
            cursor.commit()
            cursor.execute('DELETE FROM access_list WHERE username=?',(identifier))
            cursor.commit()
            flash(f'{identifier} has been deleted!')
            filter = lr_log(identifier.strip())
            lr_logLogger.addFilter(filter)
            lr_logLogger.debug(
                msg=f"doctor staff has been removed from Angel Health"
            )
            cursor.close()
        elif table == 'patient':
            cursor.execute('DELETE FROM patients WHERE username=?', (identifier))
            cursor.commit()
            cursor.execute('DELETE FROM access_list WHERE username=?',(identifier))
            cursor.commit()
            flash(f'{identifier} has been deleted!')
            filter = lr_log(identifier.strip())
            lr_logLogger.addFilter(filter)
            lr_logLogger.debug(
                msg=f"patient has been removed from Angel Health"
            )
            cursor.close()
        elif table == 'hr':
            cursor.execute('DELETE FROM hr WHERE username=?', (identifier))
            cursor.commit()
            cursor.execute('DELETE FROM access_list WHERE username=?',(identifier))
            cursor.commit()
            flash(f'{identifier} has been deleted!')
            filter = lr_log(identifier.strip())
            lr_logLogger.addFilter(filter)
            lr_logLogger.debug(
                msg=f"HR staff has been removed from Angel Health"
            )
            cursor.close()
        elif table == 'researcher':
            cursor.execute('DELETE FROM researchers WHERE username=?', (identifier))
            cursor.commit()
            cursor.execute('DELETE FROM access_list WHERE username=?',(identifier))
            cursor.commit()
            flash(f'{identifier} has been deleted!')
            filter = lr_log(identifier.strip())
            lr_logLogger.addFilter(filter)
            lr_logLogger.debug(
                msg=f"researcher staff has been removed from Angel Health"
            )
            cursor.close()
        return redirect(url_for('dashboard'))

    @app.route('/<variable>/patient', methods=['GET', 'POST'])
    def patient(variable):
        return redirect(url_for('homepage'))

    @app.route('/<variable>/doctor', methods=['GET', 'POST'])
    def doctor(variable):
        return redirect(url_for('homepage'))

    @app.route('/<variable>/admin', methods=['GET', 'POST'])
    def admin(variable):
        return redirect(url_for('homepage'))

    def shutdown_server():
        func = request.environ.get('werkzeug.server.shutdown')
        if func is None:
            raise RuntimeError('Not running with the Werkzeug Server')
        func()
        
    @app.get('/shutdown')
    def shutdown():
        shutdown_server()
        print('before')
        time.sleep(4)
        print('after')
        start_server()
       

    def start_server():
        app.run()

if __name__ == "__main__":
    add_admin()

    my_drive = MyDrive()

    vtotal = Virustotal(API_KEY="d58689de2b6f2cdec5c1625df76781dcbea39c4e705ae930da24c55f84984f40", API_VERSION="v3")
    scheduler = APScheduler()
    scheduler.init_app(app)
    scheduler.add_job(id = 'Scheduled Task',func = autonomous_conversion, trigger="interval", seconds=10)
    scheduler.start()
    app.run()

    # atexit.register(lambda: scheduler.shutdown())
